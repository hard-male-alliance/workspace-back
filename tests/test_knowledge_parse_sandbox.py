"""@brief Knowledge 文件 parser 的进程隔离回归 / Process-isolation regressions for Knowledge file parsing."""

from __future__ import annotations

import asyncio
import io
import json
import sys
from dataclasses import dataclass, field
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from backend.domain.common import DomainError
from backend.infrastructure import knowledge_parse_sandbox, knowledge_parsing
from backend.infrastructure.knowledge_parsing import (
    KnowledgeParseProcessLimits,
    LocalKnowledgeFileParser,
)
from backend.infrastructure.process_confinement import (
    ProcessConfinementMode,
    ProcessConfinementPlan,
    ProcessConfinementUnavailable,
    clear_confinement_probe_cache,
    confinement_plan_for,
)

_DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
"""@brief DOCX 标准媒体类型 / Standard DOCX media type."""


def _text_pdf(text: str) -> bytes:
    """@brief 构造含一页可提取文本的真实 PDF / Build a real one-page PDF with extractable text."""

    output = io.BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    stream = DecodedStreamObject()
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream.set_data(f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(stream)
    writer.write(output)
    return output.getvalue()


def _docx() -> bytes:
    """@brief 构造含 heading 与正文的真实 DOCX / Build a real DOCX with a heading and body."""

    document = Document()
    document.add_heading("项目经验", level=1)
    document.add_paragraph("PostgreSQL 与 pgvector 检索")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _development_parser(
    *,
    limits: KnowledgeParseProcessLimits | None = None,
    maximum_parts: int = 10_000,
    maximum_characters: int = 100_000,
) -> LocalKnowledgeFileParser:
    """@brief 构造显式 rlimit fallback 的测试 parser / Build a test parser with an explicit rlimit fallback."""

    return LocalKnowledgeFileParser(
        maximum_characters,
        maximum_input_bytes=2 * 1024 * 1024,
        maximum_parts=maximum_parts,
        process_limits=limits,
        confinement_plan=ProcessConfinementPlan(
            ProcessConfinementMode.DEVELOPMENT,
            None,
        ),
    )


def _real_strong_plan_or_skip() -> ProcessConfinementPlan:
    """@brief 获取真实 Landlock/libseccomp 计划或跳过 / Get a real Landlock/libseccomp plan or skip."""

    clear_confinement_probe_cache()
    try:
        confinement_plan_for("production")
    except ProcessConfinementUnavailable:
        pytest.skip("Landlock ABI >= 3 and libseccomp are unavailable")
    return ProcessConfinementPlan(ProcessConfinementMode.STRONG, None)


@pytest.mark.asyncio
async def test_real_isolated_interpreter_parses_txt_markdown_docx_and_pdf() -> None:
    """@brief 四种用户格式都经真实 isolated interpreter / All four user formats use a real isolated interpreter."""

    parser = _development_parser()

    plain = await parser.parse("notes.txt", "text/plain", b"alpha\n\nbeta")
    markdown = await parser.parse(
        "notes.md",
        "text/markdown",
        b"# Evidence\n\nTyped domain model",
    )
    docx = await parser.parse("evidence.docx", _DOCX_MEDIA_TYPE, _docx())
    pdf = await parser.parse(
        "evidence.pdf",
        "application/pdf",
        _text_pdf("Isolated PDF evidence"),
    )

    assert [part.text for part in plain.parts] == ["alpha", "beta"]
    assert markdown.parts[0].metadata["heading"] == "Evidence"
    assert docx.parts[0].metadata["path"] == "paragraph/2"
    assert pdf.parts[0].metadata == {"page": 1, "path": "page/1"}


@pytest.mark.asyncio
async def test_txt_docx_and_pdf_parse_under_real_strong_confinement() -> None:
    """@brief 真实 Landlock/libseccomp 下解析三类 parser / Parse three parser families under real Landlock/libseccomp."""

    parser = LocalKnowledgeFileParser(
        100_000,
        maximum_input_bytes=2 * 1024 * 1024,
        confinement_plan=_real_strong_plan_or_skip(),
    )

    assert (await parser.parse("notes.txt", "text/plain", b"strong text")).parts[0].text == (
        "strong text"
    )
    assert (await parser.parse("evidence.docx", _DOCX_MEDIA_TYPE, _docx())).parts[
        0
    ].text == "PostgreSQL 与 pgvector 检索"
    assert (
        await parser.parse(
            "evidence.pdf",
            "application/pdf",
            _text_pdf("Strong PDF"),
        )
    ).parts[0].text == "Strong PDF"


@pytest.mark.asyncio
async def test_part_amplification_is_rejected_with_a_stable_terminal_error() -> None:
    """@brief 每字符一个段落不能放大 result / One-paragraph-per-character cannot amplify the result."""

    parser = _development_parser(maximum_parts=2)

    with pytest.raises(DomainError) as raised:
        await parser.parse("notes.txt", "text/plain", b"a\n\nb\n\nc")

    assert raised.value.problem.code == "knowledge.file_too_complex"
    assert raised.value.problem.retryable is False


def test_default_result_budget_is_derived_and_below_thirty_two_mib() -> None:
    """@brief 默认 1M chars/10k parts 的 result 预算低于 32 MiB / Default result budget is below 32 MiB."""

    parser = _development_parser(maximum_characters=1_000_000)

    assert parser._maximum_result_bytes < 32 * 1024 * 1024


@pytest.mark.asyncio
async def test_parent_rejects_more_parts_than_the_closed_protocol_allows(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """@brief 即使 child 被篡改，父进程仍拒绝超量 parts / The parent rejects excess parts even from a tampered child."""

    result = {
        "ok": True,
        "parts": [
            {
                "text": value,
                "content_type": "general",
                "metadata": {"paragraph": index, "path": f"paragraph/{index}"},
            }
            for index, value in enumerate(("a", "b", "c"), start=1)
        ],
        "metadata": {"parser": "plain_text", "extracted_characters": 3},
    }
    script = (
        "import json,pathlib,sys\n"
        f"pathlib.Path(sys.argv[1]).write_text({json.dumps(result)!r}, encoding='utf-8')\n"
    )

    def fake_argv(
        result_path: Path,
        _parser_format: object,
        **_kwargs: object,
    ) -> list[str]:
        """@brief 返回伪造超量 result 的 child / Return a child forging an excessive result."""

        return [sys.executable, "-c", script, str(result_path)]

    monkeypatch.setattr(knowledge_parsing, "_process_argv", fake_argv)
    parser = _development_parser(maximum_parts=2)

    with pytest.raises(DomainError) as raised:
        await parser.parse("notes.txt", "text/plain", b"ignored")

    assert raised.value.problem.code == "knowledge.parser_unavailable"
    assert raised.value.problem.retryable is True


@pytest.mark.asyncio
async def test_parent_rejects_an_oversized_result_without_loading_it(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """@brief 父进程在 JSON decode 前拒绝超限 result / The parent rejects an oversized result before JSON decoding."""

    script = "import pathlib,sys\npathlib.Path(sys.argv[1]).write_bytes(b'x' * (64 * 1024 + 1))\n"

    def fake_argv(
        result_path: Path,
        _parser_format: object,
        **_kwargs: object,
    ) -> list[str]:
        """@brief 返回写出超限 result 的 child / Return a child writing an oversized result."""

        return [sys.executable, "-c", script, str(result_path)]

    monkeypatch.setattr(knowledge_parsing, "_process_argv", fake_argv)
    parser = _development_parser(maximum_characters=1, maximum_parts=1)
    assert parser._maximum_result_bytes == 64 * 1024

    with pytest.raises(DomainError) as raised:
        await parser.parse("notes.txt", "text/plain", b"x")

    assert raised.value.problem.code == "knowledge.parser_unavailable"


@pytest.mark.asyncio
async def test_wall_timeout_kills_parser_descendants(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """@brief wall timeout 消灭 parser 与后代进程 / A wall timeout kills the parser and its descendants."""

    marker = tmp_path / "timeout-descendant-survived"
    descendant = (
        "import pathlib,time\n"
        "time.sleep(0.25)\n"
        f"pathlib.Path({str(marker)!r}).write_text('survived', encoding='utf-8')\n"
    )
    script = (
        "import subprocess,sys,time\n"
        f"subprocess.Popen([sys.executable, '-c', {descendant!r}])\n"
        "time.sleep(30)\n"
    )

    def fake_argv(
        _result_path: Path,
        _parser_format: object,
        **_kwargs: object,
    ) -> list[str]:
        """@brief 返回会遗留后代的卡死 parser / Return a stuck parser that would leave a descendant."""

        return [sys.executable, "-c", script]

    monkeypatch.setattr(knowledge_parsing, "_process_argv", fake_argv)
    parser = _development_parser(limits=KnowledgeParseProcessLimits(wall_timeout_seconds=0.05))

    with pytest.raises(DomainError) as raised:
        await parser.parse("notes.txt", "text/plain", b"ignored")

    assert raised.value.problem.code == "knowledge.parser_timeout"
    assert raised.value.problem.retryable is True
    await asyncio.sleep(0.4)
    assert not await asyncio.to_thread(marker.exists)


@pytest.mark.asyncio
async def test_cancellation_reaps_parser_group_and_private_directory(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """@brief coroutine 取消回收进程组与私有目录 / Coroutine cancellation reaps the group and private directory."""

    ready = tmp_path / "parser-ready"
    marker = tmp_path / "cancel-descendant-survived"
    descendant = (
        "import pathlib,time\n"
        "time.sleep(0.25)\n"
        f"pathlib.Path({str(marker)!r}).write_text('survived', encoding='utf-8')\n"
    )
    script = (
        "import pathlib,subprocess,sys,time\n"
        "result_path = pathlib.Path(sys.argv[1])\n"
        f"pathlib.Path({str(ready)!r}).write_text(str(result_path.parent), encoding='utf-8')\n"
        f"subprocess.Popen([sys.executable, '-c', {descendant!r}])\n"
        "time.sleep(30)\n"
    )

    def fake_argv(
        result_path: Path,
        _parser_format: object,
        **_kwargs: object,
    ) -> list[str]:
        """@brief 返回可观测私有目录的卡死 parser / Return a stuck parser exposing its private directory."""

        return [sys.executable, "-c", script, str(result_path)]

    monkeypatch.setattr(knowledge_parsing, "_process_argv", fake_argv)
    parser = _development_parser()
    task = asyncio.create_task(parser.parse("notes.txt", "text/plain", b"ignored"))
    for _ in range(100):
        if await asyncio.to_thread(ready.is_file):
            break
        await asyncio.sleep(0.01)
    else:
        raise AssertionError("Knowledge parser did not become ready")
    private_directory = Path(await asyncio.to_thread(ready.read_text, encoding="utf-8"))

    task.cancel()
    with pytest.raises(asyncio.CancelledError):
        await task

    await asyncio.sleep(0.4)
    assert not await asyncio.to_thread(marker.exists)
    assert not await asyncio.to_thread(private_directory.exists)


def test_parser_uses_isolated_python_and_closed_child_arguments(tmp_path: Path) -> None:
    """@brief child 始终由 ``python -I -B -m`` 启动 / The child always starts with ``python -I -B -m``."""

    result_path = tmp_path / "result.json"
    result_path.touch(mode=0o600)
    argv = knowledge_parsing._process_argv(
        result_path,
        knowledge_parsing._KnowledgeFileFormat.PDF,
        maximum_input_bytes=1024,
        maximum_extracted_characters=100,
        maximum_parts=10,
        maximum_result_bytes=64 * 1024,
        limits=KnowledgeParseProcessLimits(),
        confinement_plan=ProcessConfinementPlan(
            ProcessConfinementMode.DEVELOPMENT,
            None,
        ),
    )

    assert argv[:5] == [
        sys.executable,
        "-I",
        "-B",
        "-m",
        "backend.infrastructure.knowledge_parse_sandbox",
    ]
    assert argv[6:9] == ["pdf", "development", "1024"]


def test_production_probe_occurs_when_parser_is_constructed(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """@brief production parser 在接流量前 fail-closed / A production parser fails closed before accepting traffic."""

    def unavailable(_environment: str) -> ProcessConfinementPlan:
        """@brief 模拟缺少强隔离 / Simulate missing strong confinement."""

        raise ProcessConfinementUnavailable("missing Landlock and libseccomp")

    monkeypatch.setattr(knowledge_parsing, "confinement_plan_for", unavailable)

    with pytest.raises(ProcessConfinementUnavailable, match="Landlock"):
        LocalKnowledgeFileParser(1000, deployment_environment="production")


@dataclass(slots=True)
class _FakeResourceModule:
    """@brief 记录 child hard-limit 安装顺序 / Record child hard-limit installation order."""

    RLIMIT_CORE: int = 1
    RLIMIT_CPU: int = 2
    RLIMIT_AS: int = 3
    RLIMIT_FSIZE: int = 4
    RLIMIT_NOFILE: int = 5
    RLIMIT_NPROC: int = 6
    calls: list[tuple[int, tuple[int, int]]] = field(default_factory=list)

    def setrlimit(self, resource_kind: int, limits: tuple[int, int]) -> None:
        """@brief 记录一次 soft/hard limit / Record one soft/hard limit."""

        self.calls.append((resource_kind, limits))


def test_child_applies_all_kernel_limits_before_loading_parsers() -> None:
    """@brief child 设置六类 hard limit / The child installs all six hard limits."""

    resource_module = _FakeResourceModule()
    knowledge_parse_sandbox._apply_resource_limits(
        resource_module,
        cpu_time_seconds=5,
        memory_bytes=768 * 1024 * 1024,
        result_bytes=20 * 1024 * 1024,
        open_files=32,
        processes=1,
    )

    assert resource_module.calls == [
        (resource_module.RLIMIT_CORE, (0, 0)),
        (resource_module.RLIMIT_CPU, (5, 5)),
        (resource_module.RLIMIT_AS, (768 * 1024 * 1024, 768 * 1024 * 1024)),
        (resource_module.RLIMIT_FSIZE, (20 * 1024 * 1024, 20 * 1024 * 1024)),
        (resource_module.RLIMIT_NOFILE, (32, 32)),
        (resource_module.RLIMIT_NPROC, (1, 1)),
    ]
