"""@brief Resume V2 import parser 进程隔离回归测试 / Resume V2 import-parser process-isolation regressions."""

from __future__ import annotations

import asyncio
import hashlib
import io
import sys
from collections.abc import AsyncIterator
from contextlib import asynccontextmanager
from dataclasses import dataclass, field
from pathlib import Path

import pytest
from docx import Document

from backend.application.ports.resume_worker import (
    ResumeCapabilityFailure,
    ResumeImportSource,
)
from backend.domain.principals import WorkspaceId
from backend.domain.upload_sessions import UploadSessionId
from backend.infrastructure import resume_import_sandbox, resume_worker
from backend.infrastructure.process_confinement import (
    ProcessConfinementMode,
    ProcessConfinementPlan,
    ProcessConfinementUnavailable,
    clear_confinement_probe_cache,
    confinement_plan_for,
)
from backend.infrastructure.resume_worker import (
    ResumeImportProcessLimits,
    SafeResumeImporter,
)

_WORKSPACE_ID = WorkspaceId("workspace_importsandbox1")
"""@brief 测试 Workspace / Test Workspace."""

_UPLOAD_ID = UploadSessionId("upload_importsandbox1")
"""@brief 测试 upload session / Test upload session."""


@dataclass(slots=True)
class _BytesUploadReader:
    """@brief 仅产生给定字节的 upload reader / Upload reader yielding only the given bytes."""

    payload: bytes
    """@brief 将流式返回的字节 / Bytes returned as a stream."""

    @asynccontextmanager
    async def read(
        self,
        workspace_id: WorkspaceId,
        upload_id: UploadSessionId,
    ) -> AsyncIterator[AsyncIterator[bytes]]:
        """@brief 返回一段异步 upload stream / Return a single-chunk asynchronous upload stream.

        @param workspace_id 预期 Workspace / Expected Workspace.
        @param upload_id 预期 upload ID / Expected upload ID.
        @return 一段异步字节流 / One asynchronous byte stream.
        """

        assert workspace_id == _WORKSPACE_ID
        assert upload_id == _UPLOAD_ID

        async def chunks() -> AsyncIterator[bytes]:
            """@brief 产生唯一字节块 / Yield the sole byte chunk."""

            yield self.payload

        yield chunks()


def _source(payload: bytes, media_type: str) -> ResumeImportSource:
    """@brief 为给定 payload 构造完整性证明 / Build integrity evidence for a payload.

    @param payload 已完成 upload 的字节 / Completed-upload bytes.
    @param media_type 服务端 sniff 的 MIME / Server-sniffed MIME type.
    @return worker 可信的 import source / Import source trusted by the worker.
    """

    return ResumeImportSource(
        _UPLOAD_ID,
        media_type,
        len(payload),
        hashlib.sha256(payload).hexdigest(),
    )


async def _import(
    payload: bytes,
    media_type: str,
    *,
    limits: ResumeImportProcessLimits | None = None,
    confinement_plan: ProcessConfinementPlan | None = None,
) -> str:
    """@brief 通过真实 importer 返回 plain text / Return plain text through the real importer.

    @param payload 文档字节 / Document bytes.
    @param media_type 服务端 sniff 的 MIME / Server-sniffed MIME type.
    @param limits 可选测试隔离预算 / Optional test isolation budgets.
    @param confinement_plan 可选的已探测隔离计划 / Optional pre-probed confinement plan.
    @return 解析后纯文本 / Parsed plain text.
    """

    importer = SafeResumeImporter(
        _BytesUploadReader(payload),
        process_limits=limits,
        confinement_plan=confinement_plan,
    )
    imported = await importer.import_resume(
        _WORKSPACE_ID,
        _source(payload, media_type),
        operation_id="resume.import:job_importsandbox1",
    )
    return imported.plain_text


def _real_strong_plan_or_skip() -> ProcessConfinementPlan:
    """@brief 获取当前 kernel 的真实强隔离计划，否则跳过集成测试 / Get a real strong plan or skip the integration test.

    @return 去掉可选 Bubblewrap 后的 Landlock/libseccomp 计划 / Landlock/libseccomp plan without optional Bubblewrap.
    """

    clear_confinement_probe_cache()
    try:
        plan = confinement_plan_for("production")
    except ProcessConfinementUnavailable:
        pytest.skip("Landlock ABI >= 3 and libseccomp are unavailable")
    assert plan.mode is ProcessConfinementMode.STRONG
    return ProcessConfinementPlan(ProcessConfinementMode.STRONG, None)


@pytest.mark.asyncio
async def test_import_parser_uses_a_fresh_interpreter_for_text_and_docx() -> None:
    """@brief 真实 isolated interpreter 可解析 text 与 DOCX / A real isolated interpreter parses text and DOCX."""

    assert await _import(b"Klee Example\nDistributed systems", "text/plain") == (
        "Klee Example\nDistributed systems"
    )
    document = Document()
    document.add_paragraph("Klee Example")
    document.add_paragraph("Typed domain modeling")
    buffer = io.BytesIO()
    document.save(buffer)
    assert await _import(buffer.getvalue(), resume_worker._DOCX_MEDIA_TYPE) == (
        "Klee Example\nTyped domain modeling"
    )


@pytest.mark.asyncio
async def test_import_parser_parses_text_and_docx_under_real_strong_confinement() -> None:
    """@brief TXT 与 DOCX 在真实 Landlock/libseccomp 下解析 / Parse TXT and DOCX under real Landlock/libseccomp."""

    plan = _real_strong_plan_or_skip()
    assert await _import(
        b"Klee Example\nDistributed systems",
        "text/plain",
        confinement_plan=plan,
    ) == ("Klee Example\nDistributed systems")
    document = Document()
    document.add_paragraph("Klee Example")
    document.add_paragraph("Typed domain modeling")
    buffer = io.BytesIO()
    document.save(buffer)
    assert await _import(
        buffer.getvalue(),
        resume_worker._DOCX_MEDIA_TYPE,
        confinement_plan=plan,
    ) == ("Klee Example\nTyped domain modeling")


@pytest.mark.asyncio
async def test_import_parser_returns_only_allowlisted_redacted_document_failures() -> None:
    """@brief parser exception 只映射为白名单错误码 / Parser exceptions map only to allowlisted redacted codes."""

    with pytest.raises(ResumeCapabilityFailure) as raised:
        await _import(b"not-a-pdf", "application/pdf")

    assert raised.value.code == "resume.import_invalid_document"
    assert raised.value.retryable is False
    assert str(raised.value) == "resume.import_invalid_document"


@pytest.mark.asyncio
async def test_import_parser_enforces_the_normalized_text_limit_inside_the_child() -> None:
    """@brief 超限文本不跨进程边界回传 / Oversized text is not returned across the process boundary."""

    payload = ("K" * 20_001).encode()
    with pytest.raises(ResumeCapabilityFailure) as raised:
        await _import(payload, "text/plain")

    assert raised.value.code == "resume.import_text_too_large"
    assert raised.value.retryable is False


@pytest.mark.asyncio
async def test_import_parser_hard_timeout_terminates_the_process_group(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """@brief wall timeout 会消灭 parser 及后代进程 / Wall timeout kills the parser and its descendants.

    @param tmp_path pytest 临时目录 / pytest temporary directory.
    @param monkeypatch pytest 补丁器 / pytest patch controller.
    """

    marker = tmp_path / "timeout-descendant-survived"
    descendant = (
        "import pathlib,time\n"
        "time.sleep(0.25)\n"
        f"pathlib.Path({str(marker)!r}).write_text('survived', encoding='utf-8')\n"
    )
    script = (
        "import subprocess,sys,time\n"
        f"subprocess.Popen([sys.executable, '-c', {descendant!r}])\n"
        "time.sleep(30)\n"
    )

    def fake_argv(
        _result_path: Path,
        _media_type: str,
        *,
        maximum_input_bytes: int,
        limits: ResumeImportProcessLimits,
        confinement_plan: ProcessConfinementPlan,
    ) -> list[str]:
        """@brief 返回会留下后代的卡死 parser / Return a stuck parser that leaves a descendant.

        @param _result_path 未使用 result path / Unused result path.
        @param _media_type 未使用 MIME / Unused MIME.
        @param maximum_input_bytes 输入限制 / Input limit.
        @param limits 隔离限制 / Isolation limits.
        @param confinement_plan 测试的研发隔离计划 / Test development-confinement plan.
        @return 测试 parser argv / Test parser argv.
        """

        assert maximum_input_bytes > 0
        assert limits.wall_timeout_seconds == 0.05
        assert confinement_plan.mode is ProcessConfinementMode.DEVELOPMENT
        return [sys.executable, "-c", script]

    monkeypatch.setattr(resume_worker, "_resume_import_process_argv", fake_argv)
    limits = ResumeImportProcessLimits(wall_timeout_seconds=0.05)
    with pytest.raises(ResumeCapabilityFailure) as raised:
        await _import(b"Klee", "text/plain", limits=limits)

    assert raised.value.code == "resume.import_timeout"
    assert raised.value.retryable is True
    await asyncio.sleep(0.4)
    assert not await asyncio.to_thread(marker.exists)


@pytest.mark.asyncio
async def test_import_parser_cancellation_reaps_descendants_and_temporary_directory(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """@brief coroutine 取消会回收进程组并删除私有目录 / Coroutine cancellation reaps the group and removes its private directory.

    @param tmp_path pytest 临时目录 / pytest temporary directory.
    @param monkeypatch pytest 补丁器 / pytest patch controller.
    """

    ready = tmp_path / "parser-ready"
    marker = tmp_path / "cancel-descendant-survived"
    descendant = (
        "import pathlib,time\n"
        "time.sleep(0.25)\n"
        f"pathlib.Path({str(marker)!r}).write_text('survived', encoding='utf-8')\n"
    )
    script = (
        "import pathlib,subprocess,sys,time\n"
        "result_path = pathlib.Path(sys.argv[1])\n"
        f"pathlib.Path({str(ready)!r}).write_text(str(result_path.parent), encoding='utf-8')\n"
        f"subprocess.Popen([sys.executable, '-c', {descendant!r}])\n"
        "time.sleep(30)\n"
    )

    def fake_argv(
        result_path: Path,
        _media_type: str,
        *,
        maximum_input_bytes: int,
        limits: ResumeImportProcessLimits,
        confinement_plan: ProcessConfinementPlan,
    ) -> list[str]:
        """@brief 返回可观测私有目录的卡死 parser / Return a stuck parser exposing its private directory.

        @param result_path 用于推导私有目录的路径 / Path used to infer the private directory.
        @param _media_type 未使用 MIME / Unused MIME.
        @param maximum_input_bytes 输入限制 / Input limit.
        @param limits 隔离限制 / Isolation limits.
        @param confinement_plan 测试的研发隔离计划 / Test development-confinement plan.
        @return 测试 parser argv / Test parser argv.
        """

        assert maximum_input_bytes > 0
        assert limits.wall_timeout_seconds == 10.0
        assert confinement_plan.mode is ProcessConfinementMode.DEVELOPMENT
        return [sys.executable, "-c", script, str(result_path)]

    monkeypatch.setattr(resume_worker, "_resume_import_process_argv", fake_argv)
    task = asyncio.create_task(_import(b"Klee", "text/plain"))
    for _ in range(100):
        if await asyncio.to_thread(ready.is_file):
            break
        await asyncio.sleep(0.01)
    else:
        raise AssertionError("parser child did not become ready")
    private_directory = Path(await asyncio.to_thread(ready.read_text, encoding="utf-8"))
    assert await asyncio.to_thread(private_directory.is_dir)

    task.cancel()
    with pytest.raises(asyncio.CancelledError):
        await task

    await asyncio.sleep(0.4)
    assert not await asyncio.to_thread(marker.exists)
    assert not await asyncio.to_thread(private_directory.exists)


@pytest.mark.asyncio
async def test_import_parser_rejects_an_oversized_result_before_reading_it(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """@brief 父进程在读取前拒绝超限 result / The parent rejects an oversized result before reading it.

    @param monkeypatch pytest 补丁器 / pytest patch controller.
    """

    script = (
        "import pathlib,sys\n"
        "pathlib.Path(sys.argv[1]).write_bytes(b'x' * (96 * 1024 + 1))\n"
    )

    def fake_argv(
        result_path: Path,
        _media_type: str,
        *,
        maximum_input_bytes: int,
        limits: ResumeImportProcessLimits,
        confinement_plan: ProcessConfinementPlan,
    ) -> list[str]:
        """@brief 返回会写超限 result 的 parser / Return a parser writing an oversized result.

        @param result_path 超限输出路径 / Oversized-output path.
        @param _media_type 未使用 MIME / Unused MIME.
        @param maximum_input_bytes 输入限制 / Input limit.
        @param limits 隔离限制 / Isolation limits.
        @param confinement_plan 测试的研发隔离计划 / Test development-confinement plan.
        @return 测试 parser argv / Test parser argv.
        """

        assert maximum_input_bytes > 0
        assert limits.memory_bytes > 0
        assert confinement_plan.mode is ProcessConfinementMode.DEVELOPMENT
        return [sys.executable, "-c", script, str(result_path)]

    monkeypatch.setattr(resume_worker, "_resume_import_process_argv", fake_argv)
    with pytest.raises(ResumeCapabilityFailure) as raised:
        await _import(b"Klee", "text/plain")

    assert raised.value.code == "resume.import_parser_unavailable"
    assert raised.value.retryable is True


def test_probed_bubblewrap_adds_an_optional_mount_namespace_boundary(
    tmp_path: Path,
) -> None:
    """@brief 已真实探测的 Bubblewrap 可叠加无网络最小挂载 / A probed Bubblewrap can add a no-network minimal mount layer.

    @param tmp_path pytest 临时目录 / pytest temporary directory.
    """

    result_path = tmp_path / "result.json"
    result_path.touch(mode=0o600)

    argv = resume_worker._resume_import_process_argv(
        result_path,
        "application/pdf",
        maximum_input_bytes=32 * 1024 * 1024,
        limits=ResumeImportProcessLimits(),
        confinement_plan=ProcessConfinementPlan(
            ProcessConfinementMode.STRONG,
            "/usr/bin/bwrap",
        ),
    )

    assert argv[0] == "/usr/bin/bwrap"
    assert {
        "--clearenv",
        "--die-with-parent",
        "--new-session",
        "--unshare-ipc",
        "--unshare-net",
        "--unshare-pid",
        "--unshare-user",
        "--unshare-uts",
    }.issubset(argv)
    assert ["--ro-bind", "/usr", "/usr"] == argv[
        argv.index("--ro-bind") : argv.index("--ro-bind") + 3
    ]
    result_mount = argv.index(str(result_path))
    assert argv[result_mount - 1 : result_mount + 2] == [
        "--bind",
        str(result_path),
        "/work/result.json",
    ]
    assert "/var/lib/aiws" not in argv
    assert "/var/lib/aiws-config" not in argv
    assert "--tmpfs" in argv
    assert "backend.infrastructure.resume_import_sandbox" in argv
    child_result = argv.index("backend.infrastructure.resume_import_sandbox") + 1
    assert argv[child_result] == "/work/result.json"
    assert argv[child_result + 2] == ProcessConfinementMode.STRONG.value


@dataclass(slots=True)
class _FakeResourceModule:
    """@brief 记录 setrlimit 调用的 resource 替身 / Resource substitute recording setrlimit calls."""

    RLIMIT_CORE: int = 1
    """@brief core-dump resource 标识 / Core-dump resource identifier."""

    RLIMIT_CPU: int = 2
    """@brief CPU resource 标识 / CPU resource identifier."""

    RLIMIT_AS: int = 3
    """@brief address-space resource 标识 / Address-space resource identifier."""

    RLIMIT_FSIZE: int = 4
    """@brief file-size resource 标识 / File-size resource identifier."""

    RLIMIT_NOFILE: int = 5
    """@brief descriptor resource 标识 / Descriptor resource identifier."""

    RLIMIT_NPROC: int = 6
    """@brief process resource 标识 / Process resource identifier."""

    calls: list[tuple[int, tuple[int, int]]] = field(default_factory=list)
    """@brief 已记录的 resource limits / Recorded resource limits."""

    def setrlimit(self, resource_kind: int, limits: tuple[int, int]) -> None:
        """@brief 记录一个 soft/hard limit / Record one soft/hard limit.

        @param resource_kind resource 标识 / Resource identifier.
        @param limits soft/hard 值 / Soft/hard values.
        """

        self.calls.append((resource_kind, limits))


def test_import_child_applies_every_kernel_resource_limit_before_parser_import() -> None:
    """@brief child 对六类 kernel resource 设置 hard limit / The child hard-limits all six kernel resources."""

    resource_module = _FakeResourceModule()
    resume_import_sandbox._apply_resource_limits(
        resource_module,
        cpu_time_seconds=5,
        memory_bytes=768 * 1024 * 1024,
        result_bytes=96 * 1024,
        open_files=32,
        processes=1,
    )

    assert resource_module.calls == [
        (resource_module.RLIMIT_CORE, (0, 0)),
        (resource_module.RLIMIT_CPU, (5, 5)),
        (resource_module.RLIMIT_AS, (768 * 1024 * 1024, 768 * 1024 * 1024)),
        (resource_module.RLIMIT_FSIZE, (96 * 1024, 96 * 1024)),
        (resource_module.RLIMIT_NOFILE, (32, 32)),
        (resource_module.RLIMIT_NPROC, (1, 1)),
    ]
