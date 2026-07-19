"""Parser coverage for bounded document formats and stable locators."""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from backend.domain.common import DomainError
from backend.infrastructure.knowledge_parsing import LocalKnowledgeFileParser


def _text_pdf(text: str) -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    stream = DecodedStreamObject()
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream.set_data(f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(stream)
    writer.write(output)
    return output.getvalue()


@pytest.mark.asyncio
async def test_pdf_parser_preserves_page_locator() -> None:
    parser = LocalKnowledgeFileParser(100_000)
    parsed = await parser.parse(
        "evidence.pdf",
        "application/pdf",
        _text_pdf("Kubernetes PDF evidence"),
    )
    assert parsed.parts[0].text == "Kubernetes PDF evidence"
    assert parsed.parts[0].metadata == {"page": 1, "path": "page/1"}
    assert parsed.metadata["page_count"] == 1


@pytest.mark.asyncio
async def test_docx_parser_preserves_heading_and_paragraph() -> None:
    document = Document()
    document.add_heading("项目经验", level=1)
    document.add_paragraph("PostgreSQL 与 pgvector 检索")
    output = BytesIO()
    document.save(output)
    parser = LocalKnowledgeFileParser(100_000)
    parsed = await parser.parse(
        "evidence.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        output.getvalue(),
    )
    assert parsed.parts[0].text == "PostgreSQL 与 pgvector 检索"
    assert parsed.parts[0].metadata["heading"] == "项目经验"
    assert parsed.parts[0].metadata["path"] == "paragraph/2"


@pytest.mark.asyncio
async def test_scanned_pdf_fails_closed_when_ocr_is_disabled() -> None:
    output = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.write(output)
    parser = LocalKnowledgeFileParser(100_000)
    with pytest.raises(DomainError) as raised:
        await parser.parse("scan.pdf", "application/pdf", output.getvalue())
    assert raised.value.problem.code == "knowledge.file_no_extractable_text"
