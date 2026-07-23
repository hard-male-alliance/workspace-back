"""@brief Resume worker 的真实 import/render 能力 / Real import/render capabilities for the Resume worker."""

from __future__ import annotations

import asyncio
import hashlib
import io
import json
import math
import os
import re
import signal
import stat
import sys
import tempfile
import zipfile
from collections.abc import Callable, Mapping, Sequence
from dataclasses import dataclass
from datetime import UTC
from pathlib import Path
from typing import Any, cast

from docx import Document
from pydantic import TypeAdapter
from pypdf import PdfReader

from backend.application.ports.resume_worker import (
    RenderedResumeArtifact,
    ResumeCapabilityFailure,
    ResumeImportedContent,
    ResumeImportSource,
    ResumeUploadObjectReader,
    resume_worker_artifact_id,
)
from backend.domain.common import DomainError
from backend.domain.ports import Renderer
from backend.domain.principals import WorkspaceId
from backend.domain.resume_jobs import RenderFormat
from backend.domain.resumes import JsonValue, ResumeDocument
from backend.infrastructure.process_confinement import (
    ProcessConfinementMode,
    ProcessConfinementPlan,
    confinement_plan_for,
)

_DOCUMENT_ADAPTER: TypeAdapter[ResumeDocument] = TypeAdapter(ResumeDocument)
"""@brief ResumeDocument 的规范 JSON codec / Canonical JSON codec for ResumeDocument."""

_MAXIMUM_IMPORT_BYTES = 32 * 1024 * 1024
"""@brief worker materialize 的 import 硬上限 / Hard import materialization limit for the worker."""

_MAXIMUM_PDF_PAGES = 200
"""@brief import PDF 页数硬上限 / Hard page-count limit for imported PDFs."""

_MAXIMUM_IMPORT_TEXT_CHARACTERS = 20_000
"""@brief 进入 SIR 前的文本字符上限 / Text-character limit before entering the SIR."""

_MAXIMUM_IMPORT_RESULT_BYTES = 96 * 1024
"""@brief 隔离子进程 result envelope 上限 / Isolated-child result-envelope limit."""

_IMPORT_PROCESS_TERM_GRACE_SECONDS = 0.5
"""@brief import 子进程 SIGTERM 宽限 / SIGTERM grace period for an import child."""

_IMPORT_PROCESS_KILL_REAP_SECONDS = 1.0
"""@brief SIGKILL 后的回收上限 / Reaping limit after SIGKILL."""

_IMPORT_CHILD_FAILURE_CODES = frozenset(
    {
        "resume.import_empty_document",
        "resume.import_format_unsupported",
        "resume.import_invalid_document",
        "resume.import_invalid_encoding",
        "resume.import_text_too_large",
        "resume.import_too_many_pages",
    }
)
"""@brief 子进程可回传的确定性错误白名单 / Allowlisted deterministic child errors."""

_killpg: Callable[[int, int], None] | None = getattr(os, "killpg", None)
"""@brief POSIX 进程组信号函数 / POSIX process-group signaling function."""

_DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
"""@brief DOCX 标准媒体类型 / Standard DOCX media type."""


@dataclass(frozen=True, slots=True)
class ResumeImportProcessLimits:
    """@brief Resume import 子进程的双重资源预算 / Dual resource budgets for a Resume-import child.

    @param wall_timeout_seconds 父进程强制的 wall-clock deadline / Parent-enforced wall-clock deadline.
    @param cpu_time_seconds kernel 强制的 CPU 秒数 / Kernel-enforced CPU seconds.
    @param memory_bytes 子进程 virtual address-space 上限 / Child virtual-address-space limit.
    @param open_files 子进程 file-descriptor 上限 / Child file-descriptor limit.
    @param processes 子进程可创建的进程数上限 / Child-created process limit.
    """

    wall_timeout_seconds: float = 10.0
    """@brief wall-clock 硬超时 / Hard wall-clock timeout."""

    cpu_time_seconds: int = 5
    """@brief CPU 时间硬上限 / Hard CPU-time limit."""

    memory_bytes: int = 768 * 1024 * 1024
    """@brief virtual memory 硬上限 / Hard virtual-memory limit."""

    open_files: int = 32
    """@brief 打开 descriptor 硬上限 / Hard open-descriptor limit."""

    processes: int = 1
    """@brief 可创建的子进程上限 / Child-process creation limit."""

    def __post_init__(self) -> None:
        """@brief 验证隔离预算不可被退化或无界化 / Prevent weakened or unbounded isolation budgets."""

        if (
            isinstance(self.wall_timeout_seconds, bool)
            or not math.isfinite(self.wall_timeout_seconds)
            or not 0.01 <= self.wall_timeout_seconds <= 120.0
        ):
            raise ValueError("Resume import wall timeout must be between 0.01 and 120 seconds")
        if (
            isinstance(self.cpu_time_seconds, bool)
            or not 1 <= self.cpu_time_seconds <= 60
        ):
            raise ValueError("Resume import CPU limit must be between one and 60 seconds")
        if (
            isinstance(self.memory_bytes, bool)
            or not 256 * 1024 * 1024 <= self.memory_bytes <= 2 * 1024 * 1024 * 1024
        ):
            raise ValueError("Resume import memory limit must be between 256 MiB and two GiB")
        if isinstance(self.open_files, bool) or not 16 <= self.open_files <= 256:
            raise ValueError("Resume import descriptor limit must be between 16 and 256")
        if isinstance(self.processes, bool) or not 1 <= self.processes <= 8:
            raise ValueError("Resume import process limit must be between one and eight")


class SafeResumeImporter:
    """@brief 从已安全扫描 upload 提取有界纯文本 / Extract bounded plain text from a safely scanned upload."""

    def __init__(
        self,
        reader: ResumeUploadObjectReader,
        *,
        maximum_bytes: int = _MAXIMUM_IMPORT_BYTES,
        process_limits: ResumeImportProcessLimits | None = None,
        deployment_environment: str = "development",
        confinement_plan: ProcessConfinementPlan | None = None,
    ) -> None:
        """@brief 绑定 server-side reader 与独立 worker 配额 / Bind a server-side reader and worker-specific quota.

        @param reader 只能按 Workspace/upload ID 读隔离对象的 reader / Reader restricted by Workspace and upload ID.
        @param maximum_bytes worker 内存 materialize 上限 / Worker materialization limit.
        @param process_limits 解析子进程的有界执行策略 / Bounded-execution policy for the parser child.
        @param deployment_environment 决定生产强隔离或研发 fallback / Selects production strong confinement or the development fallback.
        @param confinement_plan 测试可注入的已 probe 计划 / Test-injectable pre-probed plan.
        """
        if isinstance(maximum_bytes, bool) or not 1 <= maximum_bytes <= 128 * 1024 * 1024:
            raise ValueError("Resume import maximum bytes must be between one and 128 MiB")
        self._reader = reader
        self._maximum_bytes = maximum_bytes
        self._process_limits = process_limits or ResumeImportProcessLimits()
        self._confinement_plan = confinement_plan or confinement_plan_for(
            deployment_environment
        )

    async def import_resume(
        self,
        workspace_id: WorkspaceId,
        source: ResumeImportSource,
        *,
        operation_id: str,
    ) -> ResumeImportedContent:
        """@brief 流式重验 size/hash 后解析 PDF、DOCX 或 UTF-8 文本 / Reverify size/hash while streaming, then parse PDF, DOCX, or UTF-8 text."""
        if not operation_id:
            raise ValueError("Resume import operation ID is required")
        if source.size_bytes > self._maximum_bytes:
            raise ResumeCapabilityFailure("resume.import_too_large", retryable=False)
        payload = bytearray()
        digest = hashlib.sha256()
        try:
            async with self._reader.read(workspace_id, source.upload_session_id) as chunks:
                async for chunk in chunks:
                    if not isinstance(chunk, bytes) or not chunk:
                        continue
                    if len(payload) + len(chunk) > self._maximum_bytes:
                        raise ResumeCapabilityFailure(
                            "resume.import_too_large",
                            retryable=False,
                        )
                    payload.extend(chunk)
                    digest.update(chunk)
        except ResumeCapabilityFailure:
            raise
        except Exception as error:
            raise ResumeCapabilityFailure(
                "resume.import_source_unavailable",
                retryable=True,
            ) from error
        if len(payload) != source.size_bytes or digest.hexdigest() != source.sha256:
            raise ResumeCapabilityFailure("resume.import_integrity_mismatch", retryable=False)
        try:
            text = await _extract_import_text_in_subprocess(
                bytes(payload),
                source.media_type.casefold(),
                maximum_input_bytes=self._maximum_bytes,
                limits=self._process_limits,
                confinement_plan=self._confinement_plan,
            )
        except ResumeCapabilityFailure:
            raise
        except Exception as error:
            raise ResumeCapabilityFailure(
                "resume.import_invalid_document",
                retryable=False,
            ) from error
        normalized = _normalize_import_text(text)
        if not normalized:
            raise ResumeCapabilityFailure("resume.import_empty_document", retryable=False)
        if len(normalized) > _MAXIMUM_IMPORT_TEXT_CHARACTERS:
            raise ResumeCapabilityFailure("resume.import_text_too_large", retryable=False)
        first_line = next((line.strip("#*_- \t") for line in normalized.splitlines() if line.strip()), "")
        full_name = first_line if 1 <= len(first_line) <= 200 else "Imported candidate"
        return ResumeImportedContent(full_name, normalized)


class MultiFormatResumeRenderer:
    """@brief 复用受限 PDF renderer 并原生生成 JSON/DOCX / Reuse the restricted PDF renderer and natively generate JSON/DOCX."""

    def __init__(self, pdf_renderer: Renderer) -> None:
        """@brief 绑定部署选择的 PDF renderer / Bind the deployment-selected PDF renderer."""
        self._pdf_renderer = pdf_renderer

    async def render_resume(
        self,
        document: ResumeDocument,
        formats: Sequence[RenderFormat],
        *,
        operation_id: str,
    ) -> tuple[RenderedResumeArtifact, ...]:
        """@brief 在事务外生成所有请求格式 / Generate every requested format outside transactions."""
        if not operation_id or not formats or len(set(formats)) != len(formats):
            raise ValueError("Resume render request is invalid")
        payload = _document_payload(document)
        outputs: list[RenderedResumeArtifact] = []
        for output_format in formats:
            artifact_id = resume_worker_artifact_id(operation_id, output_format)
            if output_format is RenderFormat.PDF:
                render_payload = dict(payload)
                render_payload["artifact_id"] = str(artifact_id)
                try:
                    content, raw_source_map = await self._pdf_renderer.render(render_payload)
                except DomainError as error:
                    raise ResumeCapabilityFailure(
                        error.problem.code,
                        retryable=error.problem.retryable,
                    ) from error
                except Exception as error:
                    raise ResumeCapabilityFailure(
                        "resume.renderer_unavailable",
                        retryable=True,
                    ) from error
                source_map = _source_map(raw_source_map)
                page_count = _pdf_page_count(content)
                outputs.append(
                    RenderedResumeArtifact(
                        artifact_id,
                        RenderFormat.PDF,
                        "application/pdf",
                        content,
                        page_count,
                        source_map,
                    )
                )
            elif output_format is RenderFormat.JSON:
                content = json.dumps(
                    payload,
                    ensure_ascii=False,
                    sort_keys=True,
                    separators=(",", ":"),
                ).encode("utf-8")
                outputs.append(
                    RenderedResumeArtifact(
                        artifact_id,
                        RenderFormat.JSON,
                        "application/json",
                        content,
                    )
                )
            elif output_format is RenderFormat.DOCX:
                content = await asyncio.to_thread(_render_docx, document)
                outputs.append(
                    RenderedResumeArtifact(
                        artifact_id,
                        RenderFormat.DOCX,
                        _DOCX_MEDIA_TYPE,
                        content,
                    )
                )
            else:
                raise ResumeCapabilityFailure(
                    "resume.render_format_unsupported",
                    retryable=False,
                )
        return tuple(outputs)


async def _extract_import_text_in_subprocess(
    payload: bytes,
    media_type: str,
    *,
    maximum_input_bytes: int,
    limits: ResumeImportProcessLimits,
    confinement_plan: ProcessConfinementPlan,
) -> str:
    """@brief 在可硬终止的独立 session 中解析文档 / Parse a document in a hard-killable isolated session.

    @param payload 已在父进程验证的文档字节 / Document bytes verified by the parent.
    @param media_type 服务端 sniff 的 MIME / Server-sniffed MIME type.
    @param maximum_input_bytes 子进程 stdin 重验上限 / Child-side stdin revalidation limit.
    @param limits wall-clock 与 kernel 资源限制 / Wall-clock and kernel resource limits.
    @param confinement_plan 已完成 capability probe 的隔离计划 / Capability-probed confinement plan.
    @return 已在 child 中规范化且有界的文本 / Text normalized and bounded in the child.
    @raise ResumeCapabilityFailure 解析失败、超时或隔离进程不可用 / Parsing failure, timeout, or unavailable isolated process.
    @note stdout/stderr 均丢弃；只从 RLIMIT_FSIZE 保护的私有临时文件读取
    白名单 JSON，避免 parser 输出或 exception 反向造成父进程内存无界。
    """

    if os.name != "posix" or _killpg is None:
        raise ResumeCapabilityFailure("resume.import_parser_unavailable", retryable=True)
    with tempfile.TemporaryDirectory(prefix="aiws-resume-import-") as temporary_directory:
        workdir = Path(temporary_directory)
        result_path = workdir / "result.json"
        _create_private_import_result(result_path)
        argv = _resume_import_process_argv(
            result_path,
            media_type,
            maximum_input_bytes=maximum_input_bytes,
            limits=limits,
            confinement_plan=confinement_plan,
        )
        creation = asyncio.create_task(
            asyncio.create_subprocess_exec(
                *argv,
                cwd=workdir,
                env=_resume_import_process_environment(workdir),
                stdin=asyncio.subprocess.PIPE,
                stdout=asyncio.subprocess.DEVNULL,
                stderr=asyncio.subprocess.DEVNULL,
                start_new_session=True,
            )
        )
        try:
            process = await asyncio.shield(creation)
        except asyncio.CancelledError:
            try:
                spawned = await creation
            except (OSError, RuntimeError):
                raise
            cleanup = asyncio.create_task(_terminate_import_process_group(spawned))
            await asyncio.shield(cleanup)
            raise
        except (OSError, RuntimeError) as error:
            raise ResumeCapabilityFailure(
                "resume.import_parser_unavailable",
                retryable=True,
            ) from error
        try:
            try:
                await asyncio.wait_for(
                    process.communicate(payload),
                    timeout=limits.wall_timeout_seconds,
                )
            except TimeoutError as error:
                raise ResumeCapabilityFailure(
                    "resume.import_timeout",
                    retryable=True,
                ) from error
        finally:
            cleanup = asyncio.create_task(_terminate_import_process_group(process))
            try:
                await asyncio.shield(cleanup)
            except asyncio.CancelledError:
                await cleanup
                raise
        if process.returncode is None:
            raise ResumeCapabilityFailure("resume.import_parser_unavailable", retryable=True)
        if process.returncode < 0:
            raise ResumeCapabilityFailure("resume.import_resource_limit", retryable=False)
        if process.returncode != 0:
            raise ResumeCapabilityFailure("resume.import_parser_unavailable", retryable=True)
        return _read_import_process_result(result_path)


def _resume_import_process_argv(
    result_path: Path,
    media_type: str,
    *,
    maximum_input_bytes: int,
    limits: ResumeImportProcessLimits,
    confinement_plan: ProcessConfinementPlan,
) -> list[str]:
    """@brief 构造无 shell 且 isolated-mode 的 parser argv / Build shell-free isolated-mode parser argv.

    @param result_path 私有 result envelope 路径 / Private result-envelope path.
    @param media_type 服务端 sniff 的 MIME / Server-sniffed MIME type.
    @param maximum_input_bytes 子进程 stdin 硬上限 / Child stdin hard limit.
    @param limits 要在 child 内应用的资源限制 / Resource limits applied inside the child.
    @param confinement_plan 已 probe 的强度与可选 Bubblewrap / Probed strength and optional Bubblewrap.
    @return 可直接交给 exec 的 argv / argv ready for direct exec.
    """

    child_argv = _resume_import_child_argv(
        result_path,
        media_type,
        maximum_input_bytes=maximum_input_bytes,
        limits=limits,
        confinement_mode=confinement_plan.mode,
    )
    bubblewrap = confinement_plan.bubblewrap
    if bubblewrap is None:
        return child_argv
    return _resume_import_bubblewrap_argv(
        bubblewrap,
        result_path,
        child_argv=_resume_import_child_argv(
            Path("/work/result.json"),
            media_type,
            maximum_input_bytes=maximum_input_bytes,
            limits=limits,
            confinement_mode=confinement_plan.mode,
        ),
    )


def _resume_import_child_argv(
    result_path: Path,
    media_type: str,
    *,
    maximum_input_bytes: int,
    limits: ResumeImportProcessLimits,
    confinement_mode: ProcessConfinementMode,
) -> list[str]:
    """@brief 构造只实施 rlimit 的 child argv / Build the child argv that applies rlimits itself.

    @param result_path child 可见的 result path / Result path visible to the child.
    @param media_type 服务端 sniff 的 MIME / Server-sniffed MIME type.
    @param maximum_input_bytes 子进程 stdin 硬上限 / Child stdin hard limit.
    @param limits 子进程资源限制 / Child process resource limits.
    @param confinement_mode child 必须实施的隔离强度 / Confinement strength the child must apply.
    @return 可直接 exec 的 argv / Directly executable argv.
    """

    return [
        sys.executable,
        "-I",
        "-B",
        "-m",
        "backend.infrastructure.resume_import_sandbox",
        str(result_path),
        media_type,
        confinement_mode.value,
        str(maximum_input_bytes),
        str(_MAXIMUM_IMPORT_TEXT_CHARACTERS),
        str(_MAXIMUM_IMPORT_RESULT_BYTES),
        str(limits.cpu_time_seconds),
        str(limits.memory_bytes),
        str(limits.open_files),
        str(limits.processes),
    ]


def _resume_import_bubblewrap_argv(
    bubblewrap: str,
    result_path: Path,
    *,
    child_argv: Sequence[str],
) -> list[str]:
    """@brief 构造生产 parser 的 OS confinement argv / Build OS-confinement argv for the production parser.

    @param bubblewrap ``bwrap`` 绝对路径 / Absolute path to ``bwrap``.
    @param result_path host 上的单一可写 result 文件 / Sole writable result file on the host.
    @param child_argv 在 namespace 中执行的 parser argv / Parser argv executed inside the namespace.
    @return 无 shell 的 bubblewrap argv / Shell-free bubblewrap argv.
    @note 该边界不挂载 repository、runtime config 或数据目录；只读 Python
    runtime 与一个有界可写 result，并且移除 network namespace。
    """

    argv = [
        bubblewrap,
        "--die-with-parent",
        "--new-session",
        "--unshare-user",
        "--unshare-ipc",
        "--unshare-pid",
        "--unshare-net",
        "--unshare-uts",
        "--clearenv",
        "--setenv",
        "HOME",
        "/work",
        "--setenv",
        "LANG",
        "C.UTF-8",
        "--setenv",
        "LC_ALL",
        "C.UTF-8",
        "--setenv",
        "PATH",
        f"{sys.prefix}/bin:{os.defpath}",
        "--setenv",
        "PYTHONHASHSEED",
        "0",
        "--setenv",
        "TMPDIR",
        "/tmp",
        "--ro-bind",
        "/usr",
        "/usr",
        "--ro-bind-try",
        "/lib",
        "/lib",
        "--ro-bind-try",
        "/lib64",
        "/lib64",
        "--dir",
        "/etc",
        "--ro-bind-try",
        "/etc/ld.so.cache",
        "/etc/ld.so.cache",
        "--proc",
        "/proc",
        "--dev",
        "/dev",
        "--size",
        str(8 * 1024 * 1024),
        "--tmpfs",
        "/tmp",
        "--size",
        str(1024 * 1024),
        "--tmpfs",
        "/work",
    ]
    for runtime_root in _resume_import_runtime_roots():
        argv.extend(_bubblewrap_destination_directories(runtime_root))
        argv.extend(["--ro-bind", str(runtime_root), str(runtime_root)])
    argv.extend(
        [
            "--bind",
            str(result_path),
            "/work/result.json",
            "--chdir",
            "/work",
            *child_argv,
        ]
    )
    return argv


def _resume_import_runtime_roots() -> tuple[Path, ...]:
    """@brief 列出 ``/usr`` 之外必须只读挂载的 Python roots / List Python roots outside ``/usr`` that require read-only mounts.

    @return 去重且确定性排序的 runtime roots / Deduplicated, deterministically ordered runtime roots.
    """

    roots: set[Path] = set()
    for candidate in (Path(sys.base_prefix).resolve(), Path(sys.prefix).resolve()):
        try:
            candidate.relative_to("/usr")
        except ValueError:
            roots.add(candidate)
    return tuple(sorted(roots, key=str))


def _bubblewrap_destination_directories(path: Path) -> list[str]:
    """@brief 在空 mount namespace 中创建 runtime root 的父目录 / Create runtime-root parents in an empty mount namespace.

    @param path 将被只读挂载的绝对路径 / Absolute path to be mounted read-only.
    @return 从根到叶的 ``--dir`` argv / Root-to-leaf ``--dir`` argv.
    """

    arguments: list[str] = []
    for parent in reversed(path.parents):
        if parent == Path("/"):
            continue
        arguments.extend(["--dir", str(parent)])
    return arguments


def _create_private_import_result(result_path: Path) -> None:
    """@brief 由父进程预创建唯一可写 result inode / Have the parent pre-create the sole writable result inode.

    @param result_path 私有临时文件路径 / Private temporary file path.
    @return 无返回值 / No return value.
    """

    flags = os.O_WRONLY | os.O_CREAT | os.O_EXCL | getattr(os, "O_NOFOLLOW", 0)
    descriptor = os.open(result_path, flags, 0o600)
    os.close(descriptor)


def _resume_import_process_environment(workdir: Path) -> dict[str, str]:
    """@brief 为 parser 构造最小且私有的进程环境 / Build a minimal private process environment for the parser.

    @param workdir 只在本次调用存活的私有目录 / Private directory lasting only for this invocation.
    @return 不继承 Python/user 配置的 environment / Environment without inherited Python or user configuration.
    """

    private_directory = str(workdir)
    return {
        "HOME": private_directory,
        "LANG": "C.UTF-8",
        "LC_ALL": "C.UTF-8",
        "PATH": os.defpath,
        "PYTHONHASHSEED": "0",
        "TMPDIR": private_directory,
    }


def _read_import_process_result(result_path: Path) -> str:
    """@brief 有界读取并白名单验证 child result / Bounded-read and allowlist-validate a child result.

    @param result_path 父进程预创建、child 填充的 result path / Parent-created result path filled by the child.
    @return 已规范化的安全文本 / Normalized safe text.
    @raise ResumeCapabilityFailure result 缺失、超限、非规范或表示解析失败 / Missing, oversized, malformed, or failed result.
    """

    flags = os.O_RDONLY | getattr(os, "O_NOFOLLOW", 0)
    try:
        descriptor = os.open(result_path, flags)
    except OSError as error:
        raise ResumeCapabilityFailure(
            "resume.import_parser_unavailable",
            retryable=True,
        ) from error
    try:
        metadata = os.fstat(descriptor)
        if not stat.S_ISREG(metadata.st_mode) or metadata.st_size > _MAXIMUM_IMPORT_RESULT_BYTES:
            raise ResumeCapabilityFailure(
                "resume.import_parser_unavailable",
                retryable=True,
            )
        with os.fdopen(descriptor, "rb", closefd=False) as result_file:
            encoded = result_file.read(_MAXIMUM_IMPORT_RESULT_BYTES + 1)
    finally:
        os.close(descriptor)
    if len(encoded) > _MAXIMUM_IMPORT_RESULT_BYTES:
        raise ResumeCapabilityFailure("resume.import_parser_unavailable", retryable=True)
    try:
        value = json.loads(encoded)
    except (UnicodeDecodeError, json.JSONDecodeError) as error:
        raise ResumeCapabilityFailure(
            "resume.import_parser_unavailable",
            retryable=True,
        ) from error
    if not isinstance(value, dict) or not isinstance(value.get("ok"), bool):
        raise ResumeCapabilityFailure("resume.import_parser_unavailable", retryable=True)
    if value["ok"] is True:
        if set(value) != {"ok", "text"} or not isinstance(value.get("text"), str):
            raise ResumeCapabilityFailure("resume.import_parser_unavailable", retryable=True)
        text = cast(str, value["text"])
        if not text or len(text) > _MAXIMUM_IMPORT_TEXT_CHARACTERS:
            raise ResumeCapabilityFailure("resume.import_parser_unavailable", retryable=True)
        return text
    if set(value) != {"code", "ok"} or value.get("code") not in _IMPORT_CHILD_FAILURE_CODES:
        raise ResumeCapabilityFailure("resume.import_parser_unavailable", retryable=True)
    raise ResumeCapabilityFailure(cast(str, value["code"]), retryable=False)


async def _terminate_import_process_group(process: asyncio.subprocess.Process) -> None:
    """@brief 在所有退出路径中终止并回收 parser 进程组 / Terminate and reap the parser group on every exit path.

    @param process 独立 session 的 parser leader / Parser leader in its own session.
    @return 无返回值 / No return value.
    @note 即使 leader 已成功退出也检查进程组，防止 parser 意外留下后代。
    """

    process_group_id = process.pid
    if process.returncode is None or _import_process_group_exists(process_group_id):
        _signal_import_process_group(process_group_id, signal.SIGTERM)
    try:
        await asyncio.wait_for(process.wait(), timeout=_IMPORT_PROCESS_TERM_GRACE_SECONDS)
    except TimeoutError:
        pass
    if _import_process_group_exists(process_group_id):
        _signal_import_process_group(process_group_id, signal.SIGKILL)
    try:
        await asyncio.wait_for(process.wait(), timeout=_IMPORT_PROCESS_KILL_REAP_SECONDS)
    except TimeoutError:
        return


def _import_process_group_exists(process_group_id: int) -> bool:
    """@brief 检查 parser 进程组是否仍存在 / Check whether a parser process group still exists.

    @param process_group_id parser session leader PID/PGID / Parser session-leader PID/PGID.
    @return 进程组仍存在时为真 / True while the process group exists.
    """

    if _killpg is None:
        return False
    try:
        _killpg(process_group_id, 0)
    except ProcessLookupError:
        return False
    except PermissionError:
        return True
    return True


def _signal_import_process_group(process_group_id: int, signal_number: int) -> None:
    """@brief 向 parser 独立进程组发送信号 / Send a signal to the parser's isolated process group.

    @param process_group_id parser session leader PID/PGID / Parser session-leader PID/PGID.
    @param signal_number POSIX 信号 / POSIX signal number.
    @return 无返回值 / No return value.
    """

    if _killpg is None:
        return
    try:
        _killpg(process_group_id, signal_number)
    except (PermissionError, ProcessLookupError):
        return


def _extract_import_text(payload: bytes, media_type: str) -> str:
    """@brief 按服务端 sniff MIME 提取纯文本 / Extract plain text according to the server-sniffed MIME type."""
    if media_type in {"text/plain", "text/markdown"}:
        try:
            return payload.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise ResumeCapabilityFailure(
                "resume.import_invalid_encoding",
                retryable=False,
            ) from error
    if media_type == "application/pdf":
        reader = PdfReader(io.BytesIO(payload))
        if len(reader.pages) > _MAXIMUM_PDF_PAGES:
            raise ResumeCapabilityFailure("resume.import_too_many_pages", retryable=False)
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if media_type == _DOCX_MEDIA_TYPE:
        document = Document(io.BytesIO(payload))
        lines = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            lines.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(lines)
    raise ResumeCapabilityFailure("resume.import_format_unsupported", retryable=False)


def _normalize_import_text(value: str) -> str:
    """@brief 规范换行并移除不可持久控制字符 / Normalize newlines and remove non-persistable controls."""
    normalized = value.replace("\r\n", "\n").replace("\r", "\n")
    normalized = "".join(
        character
        for character in normalized
        if character in "\n\t" or ord(character) >= 32
    )
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in normalized.split("\n")]
    return "\n".join(lines).strip()


def _document_payload(document: ResumeDocument) -> dict[str, Any]:
    """@brief 生成 renderer 与契约 JSON 共用的 SIR payload / Build the SIR payload shared by renderers and contract JSON."""
    raw = _DOCUMENT_ADAPTER.dump_python(document, mode="json")
    if not isinstance(raw, dict):
        raise TypeError("ResumeDocument codec did not produce an object")
    payload = cast(dict[str, Any], raw)
    payload["id"] = str(document.meta.id)
    payload["revision"] = document.meta.revision
    sections = payload.get("sections")
    if isinstance(sections, list):
        for section in sections:
            if isinstance(section, dict) and isinstance(section.get("id"), str):
                section["section_id"] = section["id"]
    return payload


def _source_map(
    value: Mapping[str, Any],
) -> dict[str, JsonValue]:
    """@brief 只规范化 renderer source map，不修复其身份 / Canonicalize a renderer source map without repairing its identity.

    @param value renderer 返回的 source-map object / Source-map object returned by the renderer.
    @return 普通 canonical JSON object / Plain canonical JSON object.
    @raise ResumeCapabilityFailure 非 JSON 值立即确定性失败 / Deterministically fails for a non-JSON value.
    """
    try:
        encoded = json.dumps(dict(value), ensure_ascii=False, allow_nan=False)
        decoded = json.loads(encoded)
    except (TypeError, ValueError) as error:
        raise ResumeCapabilityFailure("resume.source_map_invalid", retryable=False) from error
    if not isinstance(decoded, dict):
        raise ResumeCapabilityFailure("resume.source_map_invalid", retryable=False)
    return cast(dict[str, JsonValue], decoded)


def _pdf_page_count(content: bytes) -> int:
    """@brief 从实际 PDF 计算可信页数 / Compute the trusted page count from the actual PDF.

    @param content renderer 返回的 PDF 内容 / PDF content returned by the renderer.
    @return 至少为一的实际页数 / Actual page count of at least one.
    @raise ResumeCapabilityFailure PDF 无效或没有页面 / Raised when the PDF is invalid or empty.
    """
    try:
        page_count = len(PdfReader(io.BytesIO(content)).pages)
    except Exception as error:
        raise ResumeCapabilityFailure("resume.render_invalid_pdf", retryable=False) from error
    if page_count < 1:
        raise ResumeCapabilityFailure("resume.render_invalid_pdf", retryable=False)
    return page_count


def _render_docx(document: ResumeDocument) -> bytes:
    """@brief 从规范 SIR 生成确定性 DOCX / Generate a deterministic DOCX from canonical SIR."""
    output = Document()
    stable_time = document.meta.created_at.astimezone(UTC).replace(tzinfo=None)
    output.core_properties.created = stable_time
    output.core_properties.modified = stable_time
    output.core_properties.author = "AI Job Workspace"
    output.add_heading(document.profile.full_name, level=0)
    if document.profile.headline:
        output.add_paragraph(document.profile.headline)
    if document.profile.summary:
        output.add_paragraph(document.profile.summary.text)
    for section in document.sections:
        if not section.visible:
            continue
        output.add_heading(section.title, level=1)
        if section.content:
            output.add_paragraph(section.content.text)
        for item in section.items:
            if not item.visible:
                continue
            heading = item.title or item.organization or item.subtitle or item.kind.value
            output.add_heading(heading, level=2)
            metadata = " · ".join(
                value
                for value in (item.organization, item.location, item.subtitle)
                if value
            )
            if metadata:
                output.add_paragraph(metadata)
            if item.summary:
                output.add_paragraph(item.summary.text)
            for highlight in item.highlights:
                output.add_paragraph(highlight.text, style="List Bullet")
    buffer = io.BytesIO()
    output.save(buffer)
    return _canonicalize_docx(buffer.getvalue())


def _canonicalize_docx(payload: bytes) -> bytes:
    """@brief 固定 ZIP entry 顺序与 timestamp 以支持 crash 重放 / Fix ZIP entry order and timestamps for crash replay."""
    source = zipfile.ZipFile(io.BytesIO(payload), "r")
    target_buffer = io.BytesIO()
    with source, zipfile.ZipFile(
        target_buffer,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as target:
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = source.getinfo(name).external_attr
            target.writestr(info, source.read(name))
    return target_buffer.getvalue()


__all__ = [
    "MultiFormatResumeRenderer",
    "ResumeImportProcessLimits",
    "SafeResumeImporter",
]
