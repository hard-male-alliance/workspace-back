"""@brief 不可信 Knowledge 文件的有界进程外解析 / Bounded out-of-process parsing for untrusted Knowledge files."""

from __future__ import annotations

import asyncio
import json
import math
import os
import re
import signal
import stat
import sys
import tempfile
from collections.abc import Callable, Mapping, Sequence
from dataclasses import dataclass
from enum import StrEnum
from io import BytesIO
from pathlib import Path
from typing import Any, cast

from backend.domain.common import DomainError, Problem
from backend.domain.knowledge import (
    KnowledgeContentType,
    KnowledgeDocumentPart,
    ParsedKnowledgeDocument,
)
from backend.infrastructure.process_confinement import (
    ProcessConfinementMode,
    ProcessConfinementPlan,
    confinement_plan_for,
)

_MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
"""@brief Markdown 标题匹配器 / Markdown-heading matcher."""

_DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
"""@brief DOCX 标准媒体类型 / Standard DOCX media type."""

_MAXIMUM_RESULT_BYTES = 64 * 1024 * 1024
"""@brief 防止配置错误造成无界 result 文件的全局硬上限 / Global hard cap preventing unbounded result files."""

_MINIMUM_RESULT_BYTES = 64 * 1024
"""@brief 小文档仍可容纳结构化 locator 的 result 下限 / Result floor accommodating locators for small documents."""

_RESULT_BYTES_PER_CHARACTER = 6
"""@brief JSON 转义后单 code point 的最大字节数 / Maximum bytes per code point after JSON escaping."""

_RESULT_BYTES_PER_PART = 2 * 1024
"""@brief 一个有界 locator 与 JSON 字段的保守预算 / Conservative budget for one bounded locator and its JSON fields."""

_MAXIMUM_LOCATOR_HEADING_CHARACTERS = 256
"""@brief 防止 heading 在每个 part metadata 中重复放大 / Prevent repeated heading amplification in part metadata."""

_PROCESS_TERM_GRACE_SECONDS = 0.5
"""@brief parser 进程组 SIGTERM 宽限 / Parser-process-group SIGTERM grace period."""

_PROCESS_KILL_REAP_SECONDS = 1.0
"""@brief SIGKILL 后的回收上限 / Reaping deadline after SIGKILL."""

_killpg: Callable[[int, int], None] | None = getattr(os, "killpg", None)
"""@brief POSIX 进程组信号函数 / POSIX process-group signaling function."""


class _KnowledgeFileFormat(StrEnum):
    """@brief 子进程协议允许的闭合文件格式 / Closed file-format set allowed by the child protocol."""

    PDF = "pdf"
    DOCX = "docx"
    TEXT = "text"
    MARKDOWN = "markdown"


@dataclass(frozen=True, slots=True)
class KnowledgeParseProcessLimits:
    """@brief Knowledge parser 的 wall-clock 与 kernel 资源预算 / Wall-clock and kernel budgets for a Knowledge parser.

    @param wall_timeout_seconds 父进程强制 wall deadline / Parent-enforced wall deadline.
    @param cpu_time_seconds kernel CPU 秒数 / Kernel CPU-second limit.
    @param memory_bytes virtual address-space 上限 / Virtual-address-space limit.
    @param open_files descriptor 上限 / Descriptor limit.
    @param processes child 可创建进程上限 / Child process limit.
    """

    wall_timeout_seconds: float = 10.0
    """@brief wall-clock 硬超时 / Hard wall-clock timeout."""

    cpu_time_seconds: int = 5
    """@brief CPU 时间硬上限 / Hard CPU-time limit."""

    memory_bytes: int = 768 * 1024 * 1024
    """@brief virtual memory 硬上限 / Hard virtual-memory limit."""

    open_files: int = 32
    """@brief 打开 descriptor 硬上限 / Hard open-descriptor limit."""

    processes: int = 1
    """@brief child 进程数上限 / Child-process limit."""

    def __post_init__(self) -> None:
        """@brief 防止隔离预算被配置为无界或不可用 / Prevent unbounded or unusable isolation budgets."""

        if (
            isinstance(self.wall_timeout_seconds, bool)
            or not math.isfinite(self.wall_timeout_seconds)
            or not 0.01 <= self.wall_timeout_seconds <= 120.0
        ):
            raise ValueError("Knowledge parser wall timeout must be between 0.01 and 120 seconds")
        if isinstance(self.cpu_time_seconds, bool) or not 1 <= self.cpu_time_seconds <= 60:
            raise ValueError("Knowledge parser CPU limit must be between one and 60 seconds")
        if (
            isinstance(self.memory_bytes, bool)
            or not 256 * 1024 * 1024 <= self.memory_bytes <= 2 * 1024 * 1024 * 1024
        ):
            raise ValueError("Knowledge parser memory limit must be between 256 MiB and two GiB")
        if isinstance(self.open_files, bool) or not 16 <= self.open_files <= 256:
            raise ValueError("Knowledge parser descriptor limit must be between 16 and 256")
        if isinstance(self.processes, bool) or not 1 <= self.processes <= 8:
            raise ValueError("Knowledge parser process limit must be between one and eight")


class LocalKnowledgeFileParser:
    """@brief 在可硬终止的隔离进程中解析用户文件 / Parse user files in a hard-killable isolated process."""

    def __init__(
        self,
        max_extracted_characters: int,
        *,
        maximum_input_bytes: int = 128 * 1024 * 1024,
        maximum_parts: int = 10_000,
        process_limits: KnowledgeParseProcessLimits | None = None,
        deployment_environment: str = "development",
        confinement_plan: ProcessConfinementPlan | None = None,
    ) -> None:
        """@brief 绑定解析预算并在部署环境启动时验证强隔离 / Bind parser budgets and probe strong confinement at deployed startup.

        @param max_extracted_characters 提取文本字符硬上限 / Hard extracted-character limit.
        @param maximum_input_bytes stdin 字节硬上限 / Hard stdin-byte limit.
        @param maximum_parts 结构化 part 数量硬上限 / Hard structured-part count limit.
        @param process_limits 独立 parser 的资源预算 / Resource budgets for each parser.
        @param deployment_environment 生产强隔离或研发 fallback 的选择 / Production-strong or development-fallback selector.
        @param confinement_plan 测试可注入的已 probe 计划 / Test-injectable pre-probed plan.
        """

        if (
            isinstance(max_extracted_characters, bool)
            or not 1 <= max_extracted_characters <= 10_000_000
        ):
            raise ValueError(
                "Knowledge parser extracted-character limit must be between one and ten million"
            )
        if (
            isinstance(maximum_input_bytes, bool)
            or not 1 <= maximum_input_bytes <= 1024 * 1024 * 1024
        ):
            raise ValueError("Knowledge parser input limit must be between one byte and one GiB")
        if isinstance(maximum_parts, bool) or not 1 <= maximum_parts <= 100_000:
            raise ValueError(
                "Knowledge parser part limit must be between one and one hundred thousand"
            )
        self._max_extracted_characters = max_extracted_characters
        self._maximum_input_bytes = maximum_input_bytes
        self._maximum_parts = maximum_parts
        self._process_limits = process_limits or KnowledgeParseProcessLimits()
        self._confinement_plan = confinement_plan or confinement_plan_for(deployment_environment)
        self._maximum_result_bytes = min(
            _MAXIMUM_RESULT_BYTES,
            max(
                _MINIMUM_RESULT_BYTES,
                max_extracted_characters * _RESULT_BYTES_PER_CHARACTER
                + maximum_parts * _RESULT_BYTES_PER_PART,
            ),
        )

    async def parse(
        self,
        filename: str,
        content_type: str,
        content: bytes,
    ) -> ParsedKnowledgeDocument:
        """@brief 选择闭合格式并通过有界进程协议解析 / Select a closed format and parse through a bounded process protocol.

        @param filename 已验证的用户文件名 / Validated user filename.
        @param content_type 服务端验证的 MIME / Server-validated MIME type.
        @param content 已完成扫描的有界 bytes / Scanned bounded bytes.
        @return 父进程重建的领域文档 / Domain document reconstructed by the parent.
        @raise DomainError 格式、内容或运行时隔离失败 / Format, content, or runtime-confinement failure.
        """

        parser_format = _parser_format_for(filename, content_type)
        if len(content) > self._maximum_input_bytes:
            raise _knowledge_problem("knowledge.file_too_large")
        return await _parse_in_subprocess(
            content,
            parser_format,
            maximum_input_bytes=self._maximum_input_bytes,
            maximum_extracted_characters=self._max_extracted_characters,
            maximum_parts=self._maximum_parts,
            maximum_result_bytes=self._maximum_result_bytes,
            limits=self._process_limits,
            confinement_plan=self._confinement_plan,
        )


def _parser_format_for(filename: str, content_type: str) -> _KnowledgeFileFormat:
    """@brief 同时校验 suffix 与 MIME 并选择 parser / Validate suffix and MIME together and select a parser.

    @param filename 文件名 / Filename.
    @param content_type 规范 MIME / Canonical MIME.
    @return 闭合 parser 格式 / Closed parser format.
    @raise DomainError suffix/MIME 组合不受支持 / Unsupported suffix/MIME combination.
    """

    suffix = Path(filename).suffix.lower()
    if content_type == "application/pdf" and suffix == ".pdf":
        return _KnowledgeFileFormat.PDF
    if content_type == _DOCX_MEDIA_TYPE and suffix == ".docx":
        return _KnowledgeFileFormat.DOCX
    if content_type in {"text/plain", "text/markdown"} and suffix in {
        ".txt",
        ".md",
        ".markdown",
    }:
        return (
            _KnowledgeFileFormat.MARKDOWN
            if content_type == "text/markdown"
            else _KnowledgeFileFormat.TEXT
        )
    raise _knowledge_problem("knowledge.file_type_unsupported")


async def _parse_in_subprocess(
    payload: bytes,
    parser_format: _KnowledgeFileFormat,
    *,
    maximum_input_bytes: int,
    maximum_extracted_characters: int,
    maximum_parts: int,
    maximum_result_bytes: int,
    limits: KnowledgeParseProcessLimits,
    confinement_plan: ProcessConfinementPlan,
) -> ParsedKnowledgeDocument:
    """@brief 在独立 session 中执行 parser 并严格验证 result / Run a parser in its own session and strictly validate its result.

    @param payload 已扫描的用户文件 / Scanned user file.
    @param parser_format 闭合格式 / Closed format.
    @param maximum_input_bytes child stdin 上限 / Child stdin limit.
    @param maximum_extracted_characters 提取文本上限 / Extracted-text limit.
    @param maximum_parts 结构化 part 上限 / Structured-part limit.
    @param maximum_result_bytes result 文件上限 / Result-file limit.
    @param limits wall-clock 与 kernel 预算 / Wall-clock and kernel budgets.
    @param confinement_plan 已 probe 的隔离计划 / Probed confinement plan.
    @return 严格重建的领域文档 / Strictly reconstructed domain document.
    @raise DomainError 超时、资源耗尽、协议错误或文档错误 / Timeout, exhaustion, protocol, or document failure.
    """

    if os.name != "posix" or _killpg is None:
        raise _knowledge_problem("knowledge.parser_unavailable")
    with tempfile.TemporaryDirectory(prefix="aiws-knowledge-parse-") as directory:
        workdir = Path(directory)
        result_path = workdir / "result.json"
        _create_private_result(result_path)
        argv = _process_argv(
            result_path,
            parser_format,
            maximum_input_bytes=maximum_input_bytes,
            maximum_extracted_characters=maximum_extracted_characters,
            maximum_parts=maximum_parts,
            maximum_result_bytes=maximum_result_bytes,
            limits=limits,
            confinement_plan=confinement_plan,
        )
        creation = asyncio.create_task(
            asyncio.create_subprocess_exec(
                *argv,
                cwd=workdir,
                env=_process_environment(workdir),
                stdin=asyncio.subprocess.PIPE,
                stdout=asyncio.subprocess.DEVNULL,
                stderr=asyncio.subprocess.DEVNULL,
                start_new_session=True,
            )
        )
        try:
            process = await asyncio.shield(creation)
        except asyncio.CancelledError:
            try:
                spawned = await creation
            except OSError, RuntimeError:
                raise
            cleanup = asyncio.create_task(_terminate_process_group(spawned))
            await asyncio.shield(cleanup)
            raise
        except (OSError, RuntimeError) as error:
            raise _knowledge_problem("knowledge.parser_unavailable") from error
        try:
            try:
                await asyncio.wait_for(
                    process.communicate(payload),
                    timeout=limits.wall_timeout_seconds,
                )
            except TimeoutError as error:
                raise _knowledge_problem("knowledge.parser_timeout") from error
        finally:
            cleanup = asyncio.create_task(_terminate_process_group(process))
            try:
                await asyncio.shield(cleanup)
            except asyncio.CancelledError:
                await cleanup
                raise
        if process.returncode is None:
            raise _knowledge_problem("knowledge.parser_unavailable")
        if process.returncode < 0:
            raise _knowledge_problem("knowledge.parser_resource_limit")
        if process.returncode != 0:
            raise _knowledge_problem("knowledge.parser_unavailable")
        return _read_process_result(
            result_path,
            parser_format,
            maximum_extracted_characters=maximum_extracted_characters,
            maximum_parts=maximum_parts,
            maximum_result_bytes=maximum_result_bytes,
        )


def _process_argv(
    result_path: Path,
    parser_format: _KnowledgeFileFormat,
    *,
    maximum_input_bytes: int,
    maximum_extracted_characters: int,
    maximum_parts: int,
    maximum_result_bytes: int,
    limits: KnowledgeParseProcessLimits,
    confinement_plan: ProcessConfinementPlan,
) -> list[str]:
    """@brief 构造无 shell、isolated-mode 的 parser argv / Build shell-free isolated-mode parser argv."""

    child = _child_argv(
        result_path,
        parser_format,
        maximum_input_bytes=maximum_input_bytes,
        maximum_extracted_characters=maximum_extracted_characters,
        maximum_parts=maximum_parts,
        maximum_result_bytes=maximum_result_bytes,
        limits=limits,
        confinement_mode=confinement_plan.mode,
    )
    if confinement_plan.bubblewrap is None:
        return child
    namespaced_child = _child_argv(
        Path("/work/result.json"),
        parser_format,
        maximum_input_bytes=maximum_input_bytes,
        maximum_extracted_characters=maximum_extracted_characters,
        maximum_parts=maximum_parts,
        maximum_result_bytes=maximum_result_bytes,
        limits=limits,
        confinement_mode=confinement_plan.mode,
    )
    return _bubblewrap_argv(
        confinement_plan.bubblewrap,
        result_path,
        child_argv=namespaced_child,
    )


def _child_argv(
    result_path: Path,
    parser_format: _KnowledgeFileFormat,
    *,
    maximum_input_bytes: int,
    maximum_extracted_characters: int,
    maximum_parts: int,
    maximum_result_bytes: int,
    limits: KnowledgeParseProcessLimits,
    confinement_mode: ProcessConfinementMode,
) -> list[str]:
    """@brief 构造 child 内部闭合协议 argv / Build argv for the closed child protocol."""

    return [
        sys.executable,
        "-I",
        "-B",
        "-m",
        "backend.infrastructure.knowledge_parse_sandbox",
        str(result_path),
        parser_format.value,
        confinement_mode.value,
        str(maximum_input_bytes),
        str(maximum_extracted_characters),
        str(maximum_parts),
        str(maximum_result_bytes),
        str(limits.cpu_time_seconds),
        str(limits.memory_bytes),
        str(limits.open_files),
        str(limits.processes),
    ]


def _bubblewrap_argv(
    bubblewrap: str,
    result_path: Path,
    *,
    child_argv: Sequence[str],
) -> list[str]:
    """@brief 叠加可选的无网络 mount namespace / Add an optional networkless mount namespace."""

    argv = [
        bubblewrap,
        "--die-with-parent",
        "--new-session",
        "--unshare-user",
        "--unshare-ipc",
        "--unshare-pid",
        "--unshare-net",
        "--unshare-uts",
        "--clearenv",
        "--setenv",
        "HOME",
        "/work",
        "--setenv",
        "LANG",
        "C.UTF-8",
        "--setenv",
        "LC_ALL",
        "C.UTF-8",
        "--setenv",
        "PATH",
        f"{sys.prefix}/bin:{os.defpath}",
        "--setenv",
        "PYTHONHASHSEED",
        "0",
        "--setenv",
        "TMPDIR",
        "/tmp",
        "--ro-bind",
        "/usr",
        "/usr",
        "--ro-bind-try",
        "/lib",
        "/lib",
        "--ro-bind-try",
        "/lib64",
        "/lib64",
        "--dir",
        "/etc",
        "--ro-bind-try",
        "/etc/ld.so.cache",
        "/etc/ld.so.cache",
        "--proc",
        "/proc",
        "--dev",
        "/dev",
        "--size",
        str(8 * 1024 * 1024),
        "--tmpfs",
        "/tmp",
        "--size",
        str(1024 * 1024),
        "--tmpfs",
        "/work",
    ]
    for runtime_root in _runtime_roots():
        argv.extend(_destination_directories(runtime_root))
        argv.extend(["--ro-bind", str(runtime_root), str(runtime_root)])
    argv.extend(
        [
            "--bind",
            str(result_path),
            "/work/result.json",
            "--chdir",
            "/work",
            *child_argv,
        ]
    )
    return argv


def _runtime_roots() -> tuple[Path, ...]:
    """@brief 列出 ``/usr`` 外的 Python runtime roots / List Python runtime roots outside ``/usr``."""

    roots: set[Path] = set()
    for candidate in (Path(sys.base_prefix).resolve(), Path(sys.prefix).resolve()):
        try:
            candidate.relative_to("/usr")
        except ValueError:
            roots.add(candidate)
    return tuple(sorted(roots, key=str))


def _destination_directories(path: Path) -> list[str]:
    """@brief 为 runtime bind mount 创建父目录 argv / Build parent-directory argv for a runtime bind mount."""

    arguments: list[str] = []
    for parent in reversed(path.parents):
        if parent != Path("/"):
            arguments.extend(["--dir", str(parent)])
    return arguments


def _create_private_result(result_path: Path) -> None:
    """@brief 由父进程预建唯一可写 result inode / Pre-create the sole writable result inode in the parent."""

    flags = os.O_WRONLY | os.O_CREAT | os.O_EXCL | getattr(os, "O_NOFOLLOW", 0)
    descriptor = os.open(result_path, flags, 0o600)
    os.close(descriptor)


def _process_environment(workdir: Path) -> dict[str, str]:
    """@brief 构造不继承 backend secrets 的最小环境 / Build a minimal environment without inherited backend secrets."""

    private_directory = str(workdir)
    return {
        "HOME": private_directory,
        "LANG": "C.UTF-8",
        "LC_ALL": "C.UTF-8",
        "PATH": os.defpath,
        "PYTHONHASHSEED": "0",
        "TMPDIR": private_directory,
    }


def _read_process_result(
    result_path: Path,
    parser_format: _KnowledgeFileFormat,
    *,
    maximum_extracted_characters: int,
    maximum_parts: int,
    maximum_result_bytes: int,
) -> ParsedKnowledgeDocument:
    """@brief 有界读取并按格式白名单重建 child result / Bounded-read and format-allowlist the child result."""

    flags = os.O_RDONLY | getattr(os, "O_NOFOLLOW", 0)
    try:
        descriptor = os.open(result_path, flags)
    except OSError as error:
        raise _knowledge_problem("knowledge.parser_unavailable") from error
    try:
        metadata = os.fstat(descriptor)
        if not stat.S_ISREG(metadata.st_mode) or metadata.st_size > maximum_result_bytes:
            raise _knowledge_problem("knowledge.parser_unavailable")
        with os.fdopen(descriptor, "rb", closefd=False) as result_file:
            encoded = result_file.read(maximum_result_bytes + 1)
    finally:
        os.close(descriptor)
    if not encoded or len(encoded) > maximum_result_bytes:
        raise _knowledge_problem("knowledge.parser_unavailable")
    try:
        value: object = json.loads(encoded)
    except (UnicodeDecodeError, json.JSONDecodeError) as error:
        raise _knowledge_problem("knowledge.parser_unavailable") from error
    if not isinstance(value, dict) or type(value.get("ok")) is not bool:
        raise _knowledge_problem("knowledge.parser_unavailable")
    if value["ok"] is False:
        if set(value) != {"code", "ok"} or not isinstance(value.get("code"), str):
            raise _knowledge_problem("knowledge.parser_unavailable")
        raise _knowledge_problem(cast(str, value["code"]), child_error=True)
    if set(value) != {"metadata", "ok", "parts"}:
        raise _knowledge_problem("knowledge.parser_unavailable")
    return _rebuild_document(
        value.get("parts"),
        value.get("metadata"),
        parser_format,
        maximum_extracted_characters=maximum_extracted_characters,
        maximum_parts=maximum_parts,
    )


def _rebuild_document(
    raw_parts: object,
    raw_metadata: object,
    parser_format: _KnowledgeFileFormat,
    *,
    maximum_extracted_characters: int,
    maximum_parts: int,
) -> ParsedKnowledgeDocument:
    """@brief 拒绝未知字段与类型并重建领域对象 / Reject unknown fields and types, then rebuild the domain object."""

    if not isinstance(raw_parts, list) or not raw_parts:
        raise _knowledge_problem("knowledge.parser_unavailable")
    if len(raw_parts) > maximum_parts:
        raise _knowledge_problem("knowledge.parser_unavailable")
    parts: list[KnowledgeDocumentPart] = []
    extracted_characters = 0
    for raw_part in raw_parts:
        if not isinstance(raw_part, dict) or set(raw_part) != {
            "content_type",
            "metadata",
            "text",
        }:
            raise _knowledge_problem("knowledge.parser_unavailable")
        text = raw_part.get("text")
        if not isinstance(text, str) or not text:
            raise _knowledge_problem("knowledge.parser_unavailable")
        if raw_part.get("content_type") != KnowledgeContentType.GENERAL.value:
            raise _knowledge_problem("knowledge.parser_unavailable")
        metadata = _validate_part_metadata(raw_part.get("metadata"), parser_format)
        extracted_characters += len(text)
        if extracted_characters > maximum_extracted_characters:
            raise _knowledge_problem("knowledge.parser_unavailable")
        parts.append(
            KnowledgeDocumentPart(
                text=text,
                content_type=KnowledgeContentType.GENERAL,
                metadata=metadata,
            )
        )
    metadata = _validate_document_metadata(
        raw_metadata,
        parser_format,
        extracted_characters=extracted_characters,
    )
    return ParsedKnowledgeDocument(tuple(parts), metadata)


def _validate_part_metadata(
    value: object,
    parser_format: _KnowledgeFileFormat,
) -> dict[str, Any]:
    """@brief 按 parser 格式验证 locator 的闭合结构 / Validate the closed locator shape for a parser format."""

    if not isinstance(value, dict):
        raise _knowledge_problem("knowledge.parser_unavailable")
    if parser_format is _KnowledgeFileFormat.PDF:
        if set(value) != {"page", "path"} or not _positive_int(value.get("page")):
            raise _knowledge_problem("knowledge.parser_unavailable")
        page = cast(int, value["page"])
        if value.get("path") != f"page/{page}":
            raise _knowledge_problem("knowledge.parser_unavailable")
    elif parser_format is _KnowledgeFileFormat.TEXT:
        if set(value) != {"paragraph", "path"} or not _positive_int(value.get("paragraph")):
            raise _knowledge_problem("knowledge.parser_unavailable")
        paragraph = cast(int, value["paragraph"])
        if value.get("path") != f"paragraph/{paragraph}":
            raise _knowledge_problem("knowledge.parser_unavailable")
    elif parser_format is _KnowledgeFileFormat.MARKDOWN:
        if set(value) != {"heading", "line_end", "line_start", "path"}:
            raise _knowledge_problem("knowledge.parser_unavailable")
        start = value.get("line_start")
        end = value.get("line_end")
        if (
            not _positive_int(start)
            or not _positive_int(end)
            or cast(int, end) < cast(int, start)
            or not _optional_string(value.get("heading"))
            or value.get("path") != f"line/{start}"
        ):
            raise _knowledge_problem("knowledge.parser_unavailable")
    else:
        keys = set(value)
        if keys == {"heading", "paragraph", "path"}:
            raw_paragraph = value.get("paragraph")
            if (
                not _positive_int(raw_paragraph)
                or not _optional_string(value.get("heading"))
                or value.get("path") != f"paragraph/{raw_paragraph}"
            ):
                raise _knowledge_problem("knowledge.parser_unavailable")
        elif keys == {"heading", "path"}:
            path = value.get("path")
            if (
                not _optional_string(value.get("heading"))
                or not isinstance(path, str)
                or re.fullmatch(r"table/[1-9]\d*/row/[1-9]\d*", path) is None
            ):
                raise _knowledge_problem("knowledge.parser_unavailable")
        else:
            raise _knowledge_problem("knowledge.parser_unavailable")
    return cast(dict[str, Any], value)


def _validate_document_metadata(
    value: object,
    parser_format: _KnowledgeFileFormat,
    *,
    extracted_characters: int,
) -> dict[str, Any]:
    """@brief 验证 parser 身份、计数和字符和 / Validate parser identity, counts, and character sum."""

    if not isinstance(value, dict) or value.get("extracted_characters") != extracted_characters:
        raise _knowledge_problem("knowledge.parser_unavailable")
    if parser_format is _KnowledgeFileFormat.PDF:
        valid = (
            set(value) == {"extracted_characters", "page_count"} | {"parser"}
            and value.get("parser") == "pypdf"
            and _positive_int(value.get("page_count"))
        )
    elif parser_format is _KnowledgeFileFormat.DOCX:
        valid = (
            set(value)
            == {
                "extracted_characters",
                "paragraph_count",
                "parser",
                "table_count",
            }
            and value.get("parser") == "python-docx"
            and _non_negative_int(value.get("paragraph_count"))
            and _non_negative_int(value.get("table_count"))
        )
    else:
        expected_parser = (
            "markdown" if parser_format is _KnowledgeFileFormat.MARKDOWN else "plain_text"
        )
        valid = (
            set(value) == {"extracted_characters", "parser"}
            and value.get("parser") == expected_parser
        )
    if not valid:
        raise _knowledge_problem("knowledge.parser_unavailable")
    return cast(dict[str, Any], value)


def _positive_int(value: object) -> bool:
    """@brief 判断值为非 bool 正整数 / Return whether a value is a non-boolean positive integer."""

    return type(value) is int and value >= 1


def _non_negative_int(value: object) -> bool:
    """@brief 判断值为非 bool 非负整数 / Return whether a value is a non-boolean non-negative integer."""

    return type(value) is int and value >= 0


def _optional_string(value: object) -> bool:
    """@brief 判断值为 string 或 None / Return whether a value is a string or ``None``."""

    return value is None or (
        isinstance(value, str) and len(value) <= _MAXIMUM_LOCATOR_HEADING_CHARACTERS
    )


async def _terminate_process_group(process: asyncio.subprocess.Process) -> None:
    """@brief 在每个退出路径终止并回收 parser 进程组 / Terminate and reap the parser process group on every exit path."""

    process_group_id = process.pid
    if process.returncode is None or _process_group_exists(process_group_id):
        _signal_process_group(process_group_id, signal.SIGTERM)
    try:
        await asyncio.wait_for(process.wait(), timeout=_PROCESS_TERM_GRACE_SECONDS)
    except TimeoutError:
        pass
    if _process_group_exists(process_group_id):
        _signal_process_group(process_group_id, signal.SIGKILL)
    try:
        await asyncio.wait_for(process.wait(), timeout=_PROCESS_KILL_REAP_SECONDS)
    except TimeoutError:
        return


def _process_group_exists(process_group_id: int) -> bool:
    """@brief 检查 parser 进程组是否仍存在 / Return whether the parser process group still exists."""

    if _killpg is None:
        return False
    try:
        _killpg(process_group_id, 0)
    except ProcessLookupError:
        return False
    except PermissionError:
        return True
    return True


def _signal_process_group(process_group_id: int, signal_number: int) -> None:
    """@brief 向独立 parser 进程组发送信号 / Send a signal to the isolated parser process group."""

    if _killpg is None:
        return
    try:
        _killpg(process_group_id, signal_number)
    except PermissionError, ProcessLookupError:
        return


_PROBLEMS: Mapping[str, Problem] = {
    "knowledge.file_type_unsupported": Problem(
        "knowledge.file_type_unsupported",
        422,
        "Knowledge file type is unsupported",
    ),
    "knowledge.file_too_large": Problem(
        "knowledge.file_too_large",
        413,
        "Knowledge file exceeds the configured size limit",
    ),
    "knowledge.file_encoding_invalid": Problem(
        "knowledge.file_encoding_invalid",
        422,
        "Text knowledge files must use UTF-8 and contain no binary NUL bytes",
    ),
    "knowledge.pdf_encrypted": Problem(
        "knowledge.pdf_encrypted",
        422,
        "Encrypted PDF files are unsupported",
    ),
    "knowledge.pdf_invalid": Problem(
        "knowledge.pdf_invalid",
        422,
        "PDF file could not be parsed",
    ),
    "knowledge.docx_invalid": Problem(
        "knowledge.docx_invalid",
        422,
        "DOCX file could not be parsed",
    ),
    "knowledge.file_no_extractable_text": Problem(
        "knowledge.file_no_extractable_text",
        422,
        "Knowledge file contains no extractable text; OCR is not enabled",
    ),
    "knowledge.extracted_text_too_large": Problem(
        "knowledge.extracted_text_too_large",
        413,
        "Extracted knowledge text exceeds the configured limit",
    ),
    "knowledge.file_too_complex": Problem(
        "knowledge.file_too_complex",
        422,
        "Knowledge file contains too many structural parts",
    ),
    "knowledge.parser_timeout": Problem(
        "knowledge.parser_timeout",
        503,
        "Knowledge file parsing timed out within the safety budget",
        retryable=True,
    ),
    "knowledge.parser_resource_limit": Problem(
        "knowledge.parser_resource_limit",
        422,
        "Knowledge file could not be parsed within the safety budget",
    ),
    "knowledge.parser_unavailable": Problem(
        "knowledge.parser_unavailable",
        503,
        "Knowledge file parser is temporarily unavailable",
        retryable=True,
    ),
}
"""@brief 跨进程可重建的公开 Problem 白名单 / Allowlist of public Problems reconstructible across the process boundary."""

_CHILD_ERROR_CODES = frozenset(
    {
        "knowledge.file_encoding_invalid",
        "knowledge.pdf_encrypted",
        "knowledge.pdf_invalid",
        "knowledge.docx_invalid",
        "knowledge.file_no_extractable_text",
        "knowledge.extracted_text_too_large",
        "knowledge.file_too_complex",
    }
)
"""@brief child 可回传的确定性文档错误 / Deterministic document errors a child may return."""


def _knowledge_problem(code: str, *, child_error: bool = False) -> DomainError:
    """@brief 从闭合 catalog 构造领域错误 / Build a domain error from the closed catalog.

    @param code 稳定错误码 / Stable error code.
    @param child_error 是否来自不可信 child envelope / Whether the code came from an untrusted child envelope.
    @return 领域错误 / Domain error.
    """

    if child_error and code not in _CHILD_ERROR_CODES:
        code = "knowledge.parser_unavailable"
    problem = _PROBLEMS.get(code, _PROBLEMS["knowledge.parser_unavailable"])
    return DomainError(problem)


def _parse_document_sync(
    payload: bytes,
    parser_format: _KnowledgeFileFormat,
    maximum_extracted_characters: int,
    maximum_parts: int,
) -> ParsedKnowledgeDocument:
    """@brief 仅供已受限 child 调用的同步 parser / Synchronous parser callable only by an already-confined child."""

    if parser_format is _KnowledgeFileFormat.PDF:
        return _parse_pdf(payload, maximum_extracted_characters, maximum_parts)
    if parser_format is _KnowledgeFileFormat.DOCX:
        return _parse_docx(payload, maximum_extracted_characters, maximum_parts)
    return _parse_text(
        payload,
        markdown=parser_format is _KnowledgeFileFormat.MARKDOWN,
        maximum_extracted_characters=maximum_extracted_characters,
        maximum_parts=maximum_parts,
    )


def _parse_text(
    content: bytes,
    *,
    markdown: bool,
    maximum_extracted_characters: int,
    maximum_parts: int,
) -> ParsedKnowledgeDocument:
    """@brief 解析 UTF-8 plain text 或 Markdown / Parse UTF-8 plain text or Markdown."""

    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise _knowledge_problem("knowledge.file_encoding_invalid") from error
    if "\x00" in text:
        raise _knowledge_problem("knowledge.file_encoding_invalid")
    parts = (
        _markdown_parts(text, maximum_parts) if markdown else _plain_text_parts(text, maximum_parts)
    )
    return _finish(
        parts,
        {"parser": "markdown" if markdown else "plain_text"},
        maximum_extracted_characters,
        maximum_parts,
    )


def _parse_pdf(
    content: bytes,
    maximum_extracted_characters: int,
    maximum_parts: int,
) -> ParsedKnowledgeDocument:
    """@brief 延迟加载 pypdf 并保留 page locator / Lazily load pypdf and preserve page locators."""

    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted and reader.decrypt("") == 0:
            raise _knowledge_problem("knowledge.pdf_encrypted")
        parts: list[KnowledgeDocumentPart] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            parts.append(
                KnowledgeDocumentPart(
                    text=text,
                    content_type=KnowledgeContentType.GENERAL,
                    metadata={"page": page_number, "path": f"page/{page_number}"},
                )
            )
            _require_part_budget(parts, maximum_parts)
    except DomainError:
        raise
    except (PdfReadError, OSError, ValueError) as error:
        raise _knowledge_problem("knowledge.pdf_invalid") from error
    return _finish(
        parts,
        {"parser": "pypdf", "page_count": len(reader.pages)},
        maximum_extracted_characters,
        maximum_parts,
    )


def _parse_docx(
    content: bytes,
    maximum_extracted_characters: int,
    maximum_parts: int,
) -> ParsedKnowledgeDocument:
    """@brief 延迟加载 python-docx 并保留 heading/paragraph locator / Lazily load python-docx and preserve locators."""

    from docx import Document

    try:
        document = Document(BytesIO(content))
    except (OSError, ValueError, KeyError) as error:
        raise _knowledge_problem("knowledge.docx_invalid") from error
    parts: list[KnowledgeDocumentPart] = []
    heading: str | None = None
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = str(paragraph.style.name or "") if paragraph.style is not None else ""
        if style_name.lower().startswith("heading"):
            heading = text[:_MAXIMUM_LOCATOR_HEADING_CHARACTERS]
            continue
        parts.append(
            KnowledgeDocumentPart(
                text=text,
                content_type=KnowledgeContentType.GENERAL,
                metadata={
                    "heading": heading,
                    "paragraph": index,
                    "path": f"paragraph/{index}",
                },
            )
        )
        _require_part_budget(parts, maximum_parts)
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                parts.append(
                    KnowledgeDocumentPart(
                        text=text,
                        content_type=KnowledgeContentType.GENERAL,
                        metadata={
                            "heading": heading,
                            "path": f"table/{table_index}/row/{row_index}",
                        },
                    )
                )
                _require_part_budget(parts, maximum_parts)
    return _finish(
        parts,
        {
            "parser": "python-docx",
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        },
        maximum_extracted_characters,
        maximum_parts,
    )


def _markdown_parts(
    text: str,
    maximum_parts: int,
) -> list[KnowledgeDocumentPart]:
    """@brief 将 Markdown 按 heading 与空行切分 / Split Markdown by headings and blank lines."""

    parts: list[KnowledgeDocumentPart] = []
    heading: str | None = None
    buffer: list[str] = []
    start_line = 1

    def flush(end_line: int) -> None:
        """@brief 将当前非空段落写入 parts / Flush the current non-empty paragraph."""

        nonlocal buffer, start_line
        value = "\n".join(buffer).strip()
        if value:
            parts.append(
                KnowledgeDocumentPart(
                    text=value,
                    content_type=KnowledgeContentType.GENERAL,
                    metadata={
                        "heading": heading,
                        "line_start": start_line,
                        "line_end": end_line,
                        "path": f"line/{start_line}",
                    },
                )
            )
            _require_part_budget(parts, maximum_parts)
        buffer = []

    lines = text.splitlines()
    for line_number, line in enumerate(lines, start=1):
        match = _MARKDOWN_HEADING.match(line)
        if match:
            flush(line_number - 1)
            heading = match.group(2).strip()[:_MAXIMUM_LOCATOR_HEADING_CHARACTERS]
            start_line = line_number + 1
        elif line.strip():
            if not buffer:
                start_line = line_number
            buffer.append(line.rstrip())
        else:
            flush(line_number - 1)
            start_line = line_number + 1
    flush(len(lines))
    return parts


def _plain_text_parts(
    text: str,
    maximum_parts: int,
) -> list[KnowledgeDocumentPart]:
    """@brief 将 plain text 按空行切分并保留段落 locator / Split plain text on blank lines with paragraph locators."""

    parts: list[KnowledgeDocumentPart] = []
    for index, paragraph in enumerate(re.split(r"\n\s*\n", text), start=1):
        value = paragraph.strip()
        if value:
            parts.append(
                KnowledgeDocumentPart(
                    text=value,
                    content_type=KnowledgeContentType.GENERAL,
                    metadata={"paragraph": index, "path": f"paragraph/{index}"},
                )
            )
            _require_part_budget(parts, maximum_parts)
    return parts


def _finish(
    parts: list[KnowledgeDocumentPart],
    metadata: dict[str, Any],
    maximum_extracted_characters: int,
    maximum_parts: int,
) -> ParsedKnowledgeDocument:
    """@brief 执行提取字符预算并完成领域文档 / Enforce the extracted-character budget and finish the domain document."""

    extracted_characters = sum(len(part.text) for part in parts)
    _require_part_budget(parts, maximum_parts)
    if extracted_characters == 0:
        raise _knowledge_problem("knowledge.file_no_extractable_text")
    if extracted_characters > maximum_extracted_characters:
        raise _knowledge_problem("knowledge.extracted_text_too_large")
    return ParsedKnowledgeDocument(
        parts=tuple(parts),
        metadata={**metadata, "extracted_characters": extracted_characters},
    )


def _require_part_budget(
    parts: Sequence[KnowledgeDocumentPart],
    maximum_parts: int,
) -> None:
    """@brief 在构建过程中尽早拒绝结构放大 / Reject structural amplification early while building."""

    if len(parts) > maximum_parts:
        raise _knowledge_problem("knowledge.file_too_complex")


__all__ = ["KnowledgeParseProcessLimits", "LocalKnowledgeFileParser"]
